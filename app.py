from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import pg8000
import os
from datetime import datetime, timedelta, timezone
import io
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
import traceback
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger
from datetime import datetime, timedelta
from werkzeug.security import generate_password_hash, check_password_hash
import hashlib
import random
import re
import json
import time
import requests
from io import BytesIO



scheduler = None

# 获取北京时间
def get_beijing_time():
    return datetime.utcnow() + timedelta(hours=8)


app = Flask(__name__)
CORS(app)

# Neon 云数据库配置
DB_CONFIG = {
    'host': 'ep-rapid-frog-ani7chkm.c-6.us-east-1.aws.neon.tech',
    'user': 'neondb_owner',
    'password': 'npg_b1QR9lMdusev',
    'database': 'neondb',
    'port': 5432
}


def get_db():
    """获取数据库连接"""
    try:
        import ssl
        ssl_context = ssl.create_default_context()

        conn = pg8000.connect(
            host=DB_CONFIG['host'],
            user=DB_CONFIG['user'],
            password=DB_CONFIG['password'],
            database=DB_CONFIG['database'],
            port=DB_CONFIG['port'],
            ssl_context=ssl_context
        )
        return conn
    except Exception as e:
        print(f"数据库连接错误: {e}")
        raise e


# ------------------- 健康检查 -------------------
@app.route("/")
def index():
    return jsonify({"status": "ok", "time": datetime.now().isoformat()})


@app.route("/health")
def health():
    return jsonify({"status": "healthy"})


@app.route("/api/test-db")
def test_db():
    """测试数据库连接"""
    try:
        db = get_db()
        cur = db.cursor()
        cur.execute("SELECT NOW()")
        result = cur.fetchone()
        cur.close()
        db.close()
        return jsonify({"status": "success", "db_time": str(result[0])})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ------------------- 初始化数据库表 -------------------
@app.route("/api/init-db", methods=["GET", "POST"])
def init_db():
    """初始化数据库表结构"""
    try:
        db = get_db()
        cur = db.cursor()

        # 创建用户表
        cur.execute("""
            CREATE TABLE IF NOT EXISTS "user" (
                id SERIAL PRIMARY KEY,
                phone VARCHAR(20) UNIQUE NOT NULL,
                password VARCHAR(100) NOT NULL,
                name VARCHAR(50),
                role VARCHAR(20) DEFAULT 'parent',
                status INTEGER DEFAULT 1,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # 创建学生表
        cur.execute("""
            CREATE TABLE IF NOT EXISTS student (
                id SERIAL PRIMARY KEY,
                name VARCHAR(50) NOT NULL,
                phone VARCHAR(20),
                grade VARCHAR(20),
                school VARCHAR(100),
                parent_name VARCHAR(50),
                parent_phone VARCHAR(20),
                status INTEGER DEFAULT 1,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # 创建教师表
        cur.execute("""
            CREATE TABLE IF NOT EXISTS teacher (
                id SERIAL PRIMARY KEY,
                name VARCHAR(50) NOT NULL,
                phone VARCHAR(20),
                subject VARCHAR(50),
                class_fee DECIMAL(10,2),
                status VARCHAR(20) DEFAULT 'active',
                hire_date DATE,
                qualification TEXT,
                bank_card VARCHAR(50),
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # 创建课时包表
        cur.execute("""
            CREATE TABLE IF NOT EXISTS course_package (
                id SERIAL PRIMARY KEY,
                student_id INTEGER REFERENCES student(id),
                total DECIMAL(8,2) NOT NULL,
                used DECIMAL(8,2) DEFAULT 0,
                surplus DECIMAL(8,2),
                course_name VARCHAR(100),
                expire_date DATE,
                status VARCHAR(20) DEFAULT 'active',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # 创建排课表
        cur.execute("""
            CREATE TABLE IF NOT EXISTS course_schedule (
                id SERIAL PRIMARY KEY,
                student_id INTEGER REFERENCES student(id),
                teacher_id INTEGER REFERENCES teacher(id),
                subject VARCHAR(100),
                class_date DATE,
                class_time VARCHAR(30),
                classroom VARCHAR(50),
                student_ids TEXT,
                duration DECIMAL(5,2) DEFAULT 2,
                status VARCHAR(20) DEFAULT 'scheduled',
                repeat_type INTEGER DEFAULT 0,
                week_day INTEGER,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # 创建课时消耗记录表
        cur.execute("""
            CREATE TABLE IF NOT EXISTS hour_consumption (
                id SERIAL PRIMARY KEY,
                package_id INTEGER REFERENCES course_package(id),
                schedule_id INTEGER REFERENCES course_schedule(id),
                hours DECIMAL(8,2) NOT NULL,
                consume_date DATE NOT NULL,
                note TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # 创建索引
        cur.execute("CREATE INDEX IF NOT EXISTS idx_cs_class_date ON course_schedule(class_date)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_cs_teacher_id ON course_schedule(teacher_id)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_cs_status ON course_schedule(status)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_hc_package_id ON hour_consumption(package_id)")

        # 插入测试数据 - 用户
        cur.execute("SELECT COUNT(*) FROM \"user\"")
        user_count = cur.fetchone()[0]
        if user_count == 0:
            cur.execute("""
                INSERT INTO "user" (phone, password, name, role) 
                VALUES ('13800138000', '123456', '管理员', 'admin')
            """)

        # 插入测试数据 - 教师
        cur.execute("SELECT COUNT(*) FROM teacher")
        teacher_count = cur.fetchone()[0]
        if teacher_count == 0:
            cur.execute("""
                INSERT INTO teacher (name, phone, subject, class_fee) 
                VALUES ('李老师', '13700137000', '数学', 150)
            """)
            cur.execute("""
                INSERT INTO teacher (name, phone, subject, class_fee) 
                VALUES ('王老师', '13800138001', '语文', 160)
            """)
            cur.execute("""
                INSERT INTO teacher (name, phone, subject, class_fee) 
                VALUES ('陈春梅', '13900139002', '数学', 180)
            """)

        # 插入测试数据 - 学生
        cur.execute("SELECT COUNT(*) FROM student")
        student_count = cur.fetchone()[0]
        if student_count == 0:
            cur.execute("""
                INSERT INTO student (name, phone, grade, school) 
                VALUES ('张三', '13900139000', '三年级', '第一小学')
            """)
            cur.execute("""
                INSERT INTO student (name, phone, grade, school) 
                VALUES ('李四', '13900139001', '四年级', '第二小学')
            """)

        db.commit()
        cur.close()
        db.close()

        return jsonify({
            "code": 200,
            "msg": "数据库初始化成功",
            "data": {
                "users": user_count,
                "teachers": teacher_count,
                "students": student_count
            }
        })
    except Exception as e:
        return jsonify({"code": 500, "msg": f"初始化失败: {str(e)}"}), 500


# ------------------- 登录接口 -------------------
from werkzeug.security import check_password_hash

@app.route("/api/login", methods=["POST"])
def login():
    try:
        data = request.get_json()
        phone = data.get('phone')
        password = data.get('password')
        code = data.get('code')
        WECHAT_APP_ID = "wx7f3bff31a3dbfd0c"
        WECHAT_APP_SECRET = "74e6b9ccbf7495205aa5e1da0a30135e"
        openid = None
        if code:
            url = f"https://api.weixin.qq.com/sns/jscode2session?appid={WECHAT_APP_ID}&secret={WECHAT_APP_SECRET}&js_code={code}&grant_type=authorization_code"
            resp = requests.get(url, timeout=5)
            wx_data = resp.json()
            openid = wx_data.get('openid')
        db = get_db()
        cur = db.cursor()
        # 查询用户时包含密码哈希字段
        cur.execute('SELECT id, name, role, openid, password FROM "user" WHERE phone = %s', (phone,))
        user = cur.fetchone()
        if user and check_password_hash(user[4], password):
            # 密码验证成功
            if openid and not user[3]:
                cur.execute('UPDATE "user" SET openid = %s WHERE id = %s', (openid, user[0]))
                db.commit()
            cur.close()
            db.close()
            return jsonify({
                "code": 200,
                "data": {
                    "id": user[0],
                    "name": user[1],
                    "role": user[2],
                    "openid": openid or user[3]
                }
            })
        else:
            cur.close()
            db.close()
            return jsonify({"code": 403, "msg": "账号或密码错误"}), 403
    except Exception as e:
        print(f"[登录] 错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500
        

# ==================== 密码修改接口（新增） ====================
@app.route("/api/user/change-password", methods=["POST"])
def change_password():
    """
    修改密码接口（基于手机号+旧密码验证）
    请求体: { "phone": "13800138000", "old_password": "xxx", "new_password": "xxx" }
    """
    try:
        data = request.get_json()
        phone = data.get('phone')
        old_password = data.get('old_password')
        new_password = data.get('new_password')
        
        if not phone or not old_password or not new_password:
            return jsonify({"code": 400, "msg": "手机号、原密码、新密码不能为空"}), 400
        
        if len(new_password) < 6:
            return jsonify({"code": 400, "msg": "新密码长度至少6位"}), 400
        
        db = get_db()
        cur = db.cursor()
        # 验证用户身份，同时获取密码哈希
        cur.execute('SELECT id, password FROM "user" WHERE phone = %s', (phone,))
        user = cur.fetchone()
        if not user:
            cur.close()
            db.close()
            return jsonify({"code": 404, "msg": "用户不存在"}), 404
        
        # 使用哈希验证原密码
        if not check_password_hash(user[1], old_password):
            cur.close()
            db.close()
            return jsonify({"code": 403, "msg": "原密码错误"}), 403
        
        # 新密码不能与旧密码相同（直接比较明文）
        if old_password == new_password:
            cur.close()
            db.close()
            return jsonify({"code": 400, "msg": "新密码不能与旧密码相同"}), 400
        
        # 生成新密码哈希并更新
        hashed_new_password = generate_password_hash(new_password)
        cur.execute('UPDATE "user" SET password = %s WHERE id = %s', (hashed_new_password, user[0]))
        db.commit()
        cur.close()
        db.close()
        
        return jsonify({"code": 200, "msg": "密码修改成功，请重新登录"})
    except Exception as e:
        print(f"修改密码错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500



def log_remind(schedule_id, remind_type, receiver_role, receiver_id, receiver_openid, 
               receiver_phone, content, status, error_msg=None, response_data=None):
    """记录提醒日志"""
    try:
        db = get_db()
        cur = db.cursor()
        cur.execute("""
            INSERT INTO remind_log 
            (schedule_id, remind_type, receiver_role, receiver_id, receiver_openid, 
             receiver_phone, content, status, error_msg, response_data, send_time)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
        """, (schedule_id, remind_type, receiver_role, receiver_id, receiver_openid,
              receiver_phone, content, status, error_msg, response_data))
        db.commit()
        cur.close()
        db.close()
        print(f"✅ 日志记录成功: {schedule_id} {status}")
    except Exception as e:
        print(f"❌ 记录提醒日志失败: {str(e)}")





# ==================== 用户管理模块（完整版）====================

@app.route("/api/admin/users", methods=["GET"])
def admin_get_users():
    """管理员获取用户列表（支持分页和搜索）"""
    try:
        page = request.args.get('page', 1, type=int)
        limit = request.args.get('limit', 20, type=int)
        keyword = request.args.get('keyword', '')
        role = request.args.get('role', '')
        
        db = get_db()
        cur = db.cursor()
        
        # 构建 WHERE 条件与参数
        where_clause = " WHERE 1=1"
        params = []
        
        if keyword:
            where_clause += " AND (u.name LIKE %s OR u.phone LIKE %s)"
            params.extend([f'%{keyword}%', f'%{keyword}%'])
        if role:
            where_clause += " AND u.role = %s"
            params.append(role)
        
        # 1. 查询总数（单独执行，避免字符串替换）
        count_sql = f"SELECT COUNT(*) FROM \"user\" u{where_clause}"
        cur.execute(count_sql, params)
        total_result = cur.fetchone()
        total = total_result[0] if total_result is not None else 0
        
        # 2. 查询分页数据
        select_sql = f"""
            SELECT 
                u.id,
                u.phone,
                u.name,
                u.role,
                u.status,
                u.created_at
            FROM \"user\" u{where_clause}
            ORDER BY u.id DESC
            LIMIT %s OFFSET %s
        """
        params.extend([limit, (page - 1) * limit])
        cur.execute(select_sql, params)
        data = cur.fetchall()
        cur.close()
        db.close()
        
        result = []
        for row in data:
            result.append({
                "id": row[0],
                "phone": row[1] or '',
                "name": row[2] or '',
                "role": row[3] or 'parent',
                "status": row[4] or 1,
                "created_at": str(row[5]) if row[5] else None
            })
        
        return jsonify({
            "code": 200,
            "data": result,
            "total": total,
            "page": page,
            "limit": limit
        })
    except Exception as e:
        print(f"获取用户列表错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500




@app.route("/api/admin/users", methods=["POST"])
def admin_add_user():
    """管理员添加用户（密码哈希存储）"""
    try:
        d = request.json
        phone = d.get('phone')
        password = d.get('password', '123456')
        name = d.get('name')
        role = d.get('role', 'parent')
        
        if not phone:
            return jsonify({"code": 400, "msg": "手机号不能为空"}), 400
        if not name:
            return jsonify({"code": 400, "msg": "姓名不能为空"}), 400
        
        db = get_db()
        cur = db.cursor()
        
        # 检查手机号是否已存在
        cur.execute("SELECT id FROM \"user\" WHERE phone = %s", (phone,))
        if cur.fetchone():
            cur.close()
            db.close()
            return jsonify({"code": 400, "msg": "手机号已存在"}), 400
        
        # 生成密码哈希
        hashed_password = generate_password_hash(password)
        
        cur.execute("""
            INSERT INTO "user" (phone, password, name, role, status)
            VALUES (%s, %s, %s, %s, 1)
            RETURNING id
        """, (phone, hashed_password, name, role))
        
        new_id = cur.fetchone()[0]
        db.commit()
        cur.close()
        db.close()
        
        return jsonify({"code": 200, "msg": "添加成功", "data": {"id": new_id}})
    except Exception as e:
        print(f"添加用户错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500

@app.route("/api/admin/users/<int:user_id>", methods=["PUT"])
def admin_update_user(user_id):
    """管理员更新用户信息"""
    try:
        d = request.json
        name = d.get('name')
        role = d.get('role')
        status = d.get('status')
        
        db = get_db()
        cur = db.cursor()
        
        updates = []
        params = []
        
        if name:
            updates.append("name = %s")
            params.append(name)
        if role:
            updates.append("role = %s")
            params.append(role)
        if status is not None:
            updates.append("status = %s")
            params.append(status)
        
        if not updates:
            return jsonify({"code": 400, "msg": "没有要更新的字段"}), 400
        
        params.append(user_id)
        sql = f"UPDATE \"user\" SET {', '.join(updates)} WHERE id = %s"
        
        cur.execute(sql, params)
        db.commit()
        cur.close()
        db.close()
        
        return jsonify({"code": 200, "msg": "更新成功"})
    except Exception as e:
        print(f"更新用户错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/admin/users/<int:user_id>", methods=["DELETE"])
def admin_delete_user(user_id):
    """管理员删除用户"""
    try:
        db = get_db()
        cur = db.cursor()
        
        cur.execute("DELETE FROM \"user\" WHERE id = %s", (user_id,))
        db.commit()
        cur.close()
        db.close()
        
        return jsonify({"code": 200, "msg": "删除成功"})
    except Exception as e:
        print(f"删除用户错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/admin/users/reset-password", methods=["POST"])
def admin_reset_password():
    """管理员重置用户密码（哈希存储）"""
    try:
        d = request.json
        user_id = d.get('user_id')
        new_password = d.get('new_password', '123456')
        
        if not user_id:
            return jsonify({"code": 400, "msg": "用户ID不能为空"}), 400
        
        db = get_db()
        cur = db.cursor()
        
        # 检查用户是否存在
        cur.execute("SELECT id FROM \"user\" WHERE id = %s", (user_id,))
        if not cur.fetchone():
            cur.close()
            db.close()
            return jsonify({"code": 404, "msg": "用户不存在"}), 404
        
        # 生成新密码的哈希值
        hashed_password = generate_password_hash(new_password)
        
        # 更新密码
        cur.execute("UPDATE \"user\" SET password = %s WHERE id = %s", (hashed_password, user_id))
        db.commit()
        cur.close()
        db.close()
        
        return jsonify({"code": 200, "msg": f"密码已重置为 {new_password}"})
    except Exception as e:
        print(f"重置密码错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500




@app.route("/api/user/list")
def user_list():
    """获取用户列表"""
    try:
        db = get_db()
        cur = db.cursor()
        cur.execute("SELECT id, phone, name, role, status, created_at FROM \"user\" ORDER BY id")
        data = cur.fetchall()
        cur.close()
        db.close()
        result = [{"id": r[0], "phone": r[1], "name": r[2], "role": r[3],
                   "status": r[4], "created_at": str(r[5]) if r[5] else None} for r in data]
        return jsonify({"code": 200, "data": result})
    except Exception as e:
        return jsonify({"code": 500, "msg": str(e)}), 500


from werkzeug.security import generate_password_hash

def migrate_passwords():
    """将明文密码转换为哈希"""
    db = get_db()
    cur = db.cursor()
    cur.execute("SELECT id, password FROM \"user\"")
    users = cur.fetchall()
    for uid, plain_pw in users:
        if plain_pw and not plain_pw.startswith('pbkdf2:sha256:'):
            hashed = generate_password_hash(plain_pw)
            cur.execute("UPDATE \"user\" SET password = %s WHERE id = %s", (hashed, uid))
    db.commit()
    cur.close()
    db.close()
    print("密码迁移完成")
    
# #临时接口
# @app.route("/migrate-passwords")
# def migrate_passwords_route():
#     migrate_passwords()
#     return "密码迁移完成"




# ==================== 系统配置模块 ====================

@app.route("/api/config/get", methods=["GET"])
def get_config():
    """获取配置（支持单个或全部）"""
    key = request.args.get('key')
    db = get_db()
    cur = db.cursor()
    if key:
        cur.execute("SELECT config_value FROM system_config WHERE config_key = %s", (key,))
        row = cur.fetchone()
        cur.close()
        db.close()
        return jsonify({"code": 200, "data": {key: row[0] if row else None}})
    else:
        cur.execute("SELECT config_key, config_value FROM system_config")
        rows = cur.fetchall()
        cur.close()
        db.close()
        return jsonify({"code": 200, "data": {r[0]: r[1] for r in rows}})


@app.route("/api/config/set", methods=["POST"])
def set_config():
    """设置配置"""
    data = request.json
    key = data.get('key')
    value = data.get('value')
    if not key:
        return jsonify({"code": 400, "msg": "缺少key"}), 400
    db = get_db()
    cur = db.cursor()
    cur.execute("""
        INSERT INTO system_config (config_key, config_value) VALUES (%s, %s)
        ON CONFLICT (config_key) DO UPDATE SET config_value = EXCLUDED.config_value, updated_at = NOW()
    """, (key, value))
    db.commit()
    cur.close()
    db.close()
    # 重启定时任务
    restart_scheduler()
    return jsonify({"code": 200, "msg": "保存成功"})




def init_scheduler():
    """初始化定时任务（从数据库读取时间）"""
    global scheduler
    if scheduler:
        scheduler.shutdown()

    scheduler = BackgroundScheduler()

    # 从数据库获取提醒时间
    db = get_db()
    cur = db.cursor()
    cur.execute("SELECT config_value FROM system_config WHERE config_key = 'remind_before_time'")
    row = cur.fetchone()
    before_time = row[0] if row else "09:00"
    cur.execute("SELECT config_value FROM system_config WHERE config_key = 'remind_after_time'")
    row = cur.fetchone()
    after_time = row[0] if row else "20:00"
    cur.close()
    db.close()

    before_hour, before_min = map(int, before_time.split(':'))
    after_hour, after_min = map(int, after_time.split(':'))

    # 课前提醒
    scheduler.add_job(
        func=scheduled_send_remind,
        trigger=CronTrigger(hour=before_hour, minute=before_min, timezone='Asia/Shanghai'),
        id='daily_remind',
        replace_existing=True
    )
    # 课后确认提醒
    scheduler.add_job(
        func=scheduled_send_confirm,
        trigger=CronTrigger(hour=after_hour, minute=after_min, timezone='Asia/Shanghai'),
        id='daily_confirm',
        replace_existing=True
    )
    scheduler.start()
    print(f"定时任务已启动：课前提醒 {before_time}，课后确认 {after_time}")


def restart_scheduler():
    """重启调度器（配置变更后调用）"""
    with app.app_context():
        init_scheduler()


@app.route("/api/remind/test-confirm", methods=["GET"])
def test_confirm():
    """手动测试课后确认提醒"""
    result = send_today_confirm_internal()
    return jsonify(result)
    
@app.route("/api/remind/test-remind", methods=["GET"])
def test_remind():
    """手动测试课前上课提醒"""
    result = send_tomorrow_remind_internal()
    return jsonify(result)



import re
import requests
import os
import time
from flask import Flask, request, jsonify, send_file
from concurrent.futures import ThreadPoolExecutor
import uuid

app = Flask(__name__)

# B站爬虫配置
BILI_COOKIES = {
    'buvid3': 'E1AC41CE-8298-B4E8-2B11-711D83FCEB4D07978infoc',
    'b_nut': '1774395407',
    'bsource': 'search_bing',
    '_uuid': '696B6B59-DCBB-93710-106A9-2F84CAC10BCA957649infoc',
    'home_feed_column': '5',
    'browser_resolution': '1912-956',
    'buvid4': '4DB22FA4-9C1B-3FF9-09D1-FC56BA6029AA09187-026032507-s5Db6HTmRJMFxI7gzVHFJA%3D%3D',
    'buvid_fp': '82d07d9422c9ad7c67f3c53cc409b7e8',
    'bmg_af_switch': '1',
    'bmg_src_def_domain': 'i2.hdslb.com',
    'bili_ticket': 'eyJhbGciOiJIUzI1NiIsImtpZCI6InMwMyIsInR5cCI6IkpXVCJ9.eyJleHAiOjE3NzQ2NTQ2MjQsImlhdCI6MTc3NDM5NTM2NCwicGx0IjotMX0.KbRAOmZjhvBQ7MrPSaVnGd_5qhwK_WyiYs25aY5WJ5I',
    'bili_ticket_expires': '1774654564',
    'sid': '83jnmdei',
    'CURRENT_QUALITY': '0',
    'rpdid': "|(umR~lRuRlJ0J'u~~RJmlRRu",
    'CURRENT_FNVAL': '4048',
    'b_lsid': 'C9BCD791_19D225F6A76',
}

BILI_HEADERS = {
    'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
    'accept-language': 'zh-CN,zh;q=0.9,en;q=0.8',
    'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/144.0.0.0 Safari/537.36 Edg/144.0.0.0',
    'referer': 'https://search.bilibili.com/',
}


@app.route('/api/video/search', methods=['POST'])
def search_videos():
    """搜索B站视频"""
    data = request.get_json()
    keyword = data.get('keyword', '')
    if not keyword:
        return jsonify({'code': 400, 'msg': '请输入搜索关键词'})
    
    params = {
        'keyword': keyword,
        'from_source': 'webtop_search',
        'spm_id_from': '333.1007',
        'search_source': '5',
    }
    try:
        response = requests.get('https://search.bilibili.com/all/', params=params, cookies=BILI_COOKIES, headers=BILI_HEADERS, timeout=10)
        urls = re.findall(r'<a href="(.*?)"', response.text)
        video_list = []
        seen = set()
        for url in urls:
            if "video" in url and url not in seen:
                full_url = 'https:' + url if url.startswith('//') else url
                seen.add(url)
                video_list.append(full_url)
                if len(video_list) >= 20:
                    break
        return jsonify({'code': 200, 'data': video_list, 'count': len(video_list)})
    except Exception as e:
        return jsonify({'code': 500, 'msg': str(e)})


import subprocess
import os
import uuid
import requests
import re
from flask import request, jsonify, send_file



@app.route('/api/video/download/simple', methods=['POST'])
def download_video_simple():
    """仅下载视频，不合成音频"""
    data = request.get_json()
    video_url = data.get('url', '')
    if not video_url:
        return jsonify({'code': 400, 'msg': '视频链接不能为空'})
    video_path = None
    audio_path = None
    output_path = None
    try:
        resp = requests.get(video_url, cookies=BILI_COOKIES, headers=BILI_HEADERS, timeout=10)
        html = resp.text
        
        # 提取标题
        title_match = re.search(r'<title>(.*?)_哔哩哔哩_bilibili</title>', html)
        title = title_match.group(1).replace(' ', '').replace('|', '').replace("'", '').replace('/', '_') if title_match else 'video'
        
        # 提取视频地址
        video_match = re.search(r'"video".*?"baseUrl":"(.*?)"', html)
        if not video_match:
            return jsonify({'code': 404, 'msg': '未找到视频源'})
        
        video_url_real = video_match.group(1)
        
        # 流式下载并直接返回
        def generate():
            response = requests.get(video_url_real, headers=BILI_HEADERS, stream=True, timeout=60)
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    yield chunk
        
        return Response(generate(), mimetype='video/mp4', headers={
            'Content-Disposition': f'attachment; filename={title}.mp4'
        })
    except Exception as e:
        return jsonify({'code': 500, 'msg': str(e)})


@app.route('/api/video/test', methods=['POST'])
def video_test():
    """使用 B站官方 API 获取视频信息（无需 Cookie）"""
    data = request.get_json()
    video_url = data.get('url', '')
    if not video_url:
        return jsonify({'code': 400, 'msg': '链接为空'})
    
    # 从 URL 中提取 BV 号
    bv_match = re.search(r'BV([a-zA-Z0-9]+)', video_url)
    if not bv_match:
        return jsonify({'code': 400, 'msg': '未找到 BV 号'})
    bv_id = bv_match.group(0)
    
    try:
        # 1. 获取视频基本信息
        api_url = f'https://api.bilibili.com/x/web-interface/view?bvid={bv_id}'
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        resp = requests.get(api_url, headers=headers, timeout=10)
        data = resp.json()
        
        if data.get('code') != 0:
            return jsonify({'code': 404, 'msg': f'API 返回错误: {data.get("message")}'})
        
        video_data = data['data']
        title = video_data.get('title', '未知标题')
        cid = video_data.get('cid')
        aid = video_data.get('aid')
        
        # 2. 获取视频播放地址
        play_url = f'https://api.bilibili.com/x/player/playurl?bvid={bv_id}&cid={cid}&qn=80'
        play_resp = requests.get(play_url, headers=headers, timeout=10)
        play_data = play_resp.json()
        
        video_url_real = None
        audio_url_real = None
        if play_data.get('code') == 0:
            dash = play_data['data'].get('dash', {})
            if dash.get('video'):
                video_url_real = dash['video'][0].get('baseUrl')
            if dash.get('audio'):
                audio_url_real = dash['audio'][0].get('baseUrl')
        
        return jsonify({
            'code': 200,
            'data': {
                'title': title,
                'bv_id': bv_id,
                'cid': cid,
                'aid': aid,
                'video_found': bool(video_url_real),
                'audio_found': bool(audio_url_real),
                'video_url': video_url_real,
                'audio_url': audio_url_real,
                'source': 'bilibili_api'
            }
        })
    except Exception as e:
        return jsonify({'code': 500, 'msg': str(e)})



@app.route('/api/video/download', methods=['POST'])
def download_video():
    """使用 B站官方 API 下载视频（无需 Cookie）"""
    data = request.get_json()
    video_url = data.get('url', '')
    if not video_url:
        return jsonify({'code': 400, 'msg': '视频链接不能为空'})
    
    # 提取 BV 号
    bv_match = re.search(r'BV([a-zA-Z0-9]+)', video_url)
    if not bv_match:
        return jsonify({'code': 400, 'msg': '未找到 BV 号'})
    bv_id = bv_match.group(0)
    
    try:
        # 1. 获取视频基本信息
        api_url = f'https://api.bilibili.com/x/web-interface/view?bvid={bv_id}'
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        resp = requests.get(api_url, headers=headers, timeout=10)
        data = resp.json()
        
        if data.get('code') != 0:
            return jsonify({'code': 404, 'msg': f'API 返回错误: {data.get("message")}'})
        
        video_data = data['data']
        title = video_data.get('title', '未知标题').replace(' ', '').replace('|', '').replace("'", '').replace('/', '_')
        cid = video_data.get('cid')
        
        # 2. 获取播放地址
        play_url = f'https://api.bilibili.com/x/player/playurl?bvid={bv_id}&cid={cid}&qn=80'
        play_resp = requests.get(play_url, headers=headers, timeout=10)
        play_data = play_resp.json()
        
        if play_data.get('code') != 0:
            return jsonify({'code': 404, 'msg': '获取播放地址失败'})
        
        dash = play_data['data'].get('dash', {})
        video_url_real = dash['video'][0]['baseUrl'] if dash.get('video') else None
        audio_url_real = dash['audio'][0]['baseUrl'] if dash.get('audio') else None
        
        if not video_url_real:
            return jsonify({'code': 404, 'msg': '未找到视频源'})
        
        # 3. 检查 FFmpeg 是否可用
        ffmpeg_available = False
        try:
            subprocess.run(['ffmpeg', '-version'], capture_output=True, timeout=5, check=False)
            ffmpeg_available = True
        except:
            pass
        
        temp_id = str(uuid.uuid4())[:8]
        video_path = f'/tmp/{temp_id}_{title}.mp4'
        audio_path = f'/tmp/{temp_id}_{title}.mp3'
        output_path = f'/tmp/{temp_id}_{title}_合成.mp4'
        
        # 4. 下载视频
        print(f"开始下载视频: {title}")
        v_resp = requests.get(video_url_real, headers=headers, timeout=60, stream=True)
        with open(video_path, 'wb') as f:
            for chunk in v_resp.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
        
        # 5. 如果有音频且 FFmpeg 可用，下载音频并合成
        if audio_url_real and ffmpeg_available:
            try:
                print("开始下载音频...")
                a_resp = requests.get(audio_url_real, headers=headers, timeout=60, stream=True)
                with open(audio_path, 'wb') as f:
                    for chunk in a_resp.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)
                
                print("开始合成视频...")
                cmd = f'ffmpeg -i "{video_path}" -i "{audio_path}" -c:v copy -c:a copy "{output_path}" -y -loglevel error'
                result = subprocess.run(cmd, shell=True, timeout=300, capture_output=True)
                
                if result.returncode == 0 and os.path.exists(output_path):
                    # 返回合成文件
                    return send_file(
                        output_path,
                        as_attachment=True,
                        download_name=f'{title}.mp4',
                        mimetype='video/mp4'
                    )
            except Exception as e:
                print(f"合成失败: {e}")
        
        # 6. 如果没有音频或合成失败，直接返回视频文件
        return send_file(
            video_path,
            as_attachment=True,
            download_name=f'{title}.mp4',
            mimetype='video/mp4'
        )
        
    except Exception as e:
        print(f"下载错误: {e}")
        return jsonify({'code': 500, 'msg': str(e)})
    finally:
        # 清理临时文件
        for path in [video_path, audio_path, output_path]:
            if path and os.path.exists(path):
                try:
                    os.remove(path)
                except:
                    pass


@app.route('/')
def index():
    return "服务运行正常"





# ==================== 文档翻译模块 ====================
# 有道翻译相关函数
def md5_hex(text):
    return hashlib.md5(text.encode('utf-8')).hexdigest()

def generate_sign(params, secret_key):
    i = params.copy()
    keys_to_delete = [k for k, v in i.items() if v == ""]
    for k in keys_to_delete:
        del i[k]
    o = sorted([k for k in i.keys() if i[k] is not None])
    o.append("key")
    i["key"] = secret_key
    a = "&".join([f"{k}={i[k]}" for k in o])
    sign = md5_hex(a)
    point_param = ",".join(o[:-1])
    return sign, point_param

def get_yduuid():
    random_num = str(random.randint(100, 9999))
    return md5_hex(random_num)

def extract_translation_from_sse(response_text):
    pattern = r'data:{"model":"yd_agent","content":"(.*?)"}'
    matches = re.findall(pattern, response_text)
    if matches:
        return "".join(matches)
    pattern2 = r'"content":"([^"]*)"'
    matches2 = re.findall(pattern2, response_text)
    if matches2:
        return "".join([m for m in matches2 if m])
    return None

def extract_by_line_parsing(response_text):
    translation_parts = []
    lines = response_text.strip().split('\n')
    for line in lines:
        if line.startswith('data:'):
            data_str = line[5:]
            try:
                data = json.loads(data_str)
                if 'content' in data:
                    translation_parts.append(data['content'])
            except:
                match = re.search(r'"content":"(.*?)"', data_str)
                if match:
                    translation_parts.append(match.group(1))
    return "".join(translation_parts)

def _single_translate(text, secret_key="QGCmKoX7N7wACtICx3PSaWzoPJza2DSC"):
    if not text:
        return ""
    yduuid = get_yduuid()
    headers = {
        'Referer': 'https://fanyi.youdao.com/',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    }
    token_url = f'https://luna-ai.youdao.com/translate_llm/secret?keyid=ai-translate-llm-pre&sign=dd26324657a782cf42f707647fa67d08&client=fanyideskweb&product=webfanyi&appVersion=12.0.0&vendor=web&pointParam=client,mysticTime,product&mysticTime=1776899320882&keyfrom=fanyi.web&mid=1&screen=1&model=1&network=wifi&abtest=0&yduuid={yduuid}'
    res = requests.get(url=token_url, headers=headers)
    token = res.json()['data']['token']
    time_stamp = int(time.time() * 1000)

    params_for_sign = {
        "appVersion": "12.0.0",
        "client": "webaitrans",
        "free": "false",
        "fromLang": "auto",
        "functionEnglishName": "LLM_translate",
        "imei": "1",
        "input": text,
        "keyfrom": "webfanyi.webaitrans",
        "keyid": "ai-translate-llm",
        "mid": "1",
        "model": "1",
        "mysticTime": str(time_stamp),
        "network": "wifi",
        "product": "webfanyi",
        "roundNo": "1",
        "screen": "1",
        "showSuggest": "0",
        "singleBox": "false",
        "source": "webaitrans",
        "token": token,
        "useTerm": "0",
        "vendor": "web",
        "yduuid": yduuid
    }
    sign, point_param = generate_sign(params_for_sign, secret_key)
    post_data = {**params_for_sign, "sign": sign, "pointParam": point_param}
    cookies = {
        'OUTFOX_SEARCH_USER_ID': '-1116170267@113.111.81.100',
        'OUTFOX_SEARCH_USER_ID_NCOO': '31318005.898740534',
    }
    headers_post = {
        'Accept': '*/*',
        'Accept-Language': 'zh-CN,zh;q=0.9',
        'Content-Type': 'application/x-www-form-urlencoded',
        'Origin': 'https://fanyi.youdao.com',
        'Referer': 'https://fanyi.youdao.com/',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
    }
    response = requests.post(
        'https://luna-ai.youdao.com/translate_llm/v3/chat',
        cookies=cookies,
        headers=headers_post,
        data=post_data,
        timeout=30
    )
    translation = extract_translation_from_sse(response.text)
    if not translation:
        translation = extract_by_line_parsing(response.text)
    return translation

def translate_text(text):
    """翻译长文本（自动分段）"""
    if not text:
        return ""
    max_len = 5000
    if len(text) > max_len:
        paragraphs = text.split('\n\n')
        chunks = []
        current = ""
        for para in paragraphs:
            if len(current) + len(para) + 2 <= max_len:
                if current:
                    current += "\n\n" + para
                else:
                    current = para
            else:
                if current:
                    chunks.append(current)
                current = para
        if current:
            chunks.append(current)
        results = []
        for chunk in chunks:
            trans = _single_translate(chunk)
            if trans:
                results.append(trans)
            time.sleep(1)  # 避免请求过快
        return "\n\n".join(results)
    else:
        return _single_translate(text)

def read_word_file(file_bytes):
    """从字节流读取 docx 内容"""
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except ImportError:
        return None

# ========== 翻译路由 ==========
@app.route("/api/translate/text", methods=["POST"])
def translate_text_api():
    try:
        data = request.get_json()
        text = data.get('text', '')
        if not text:
            return jsonify({"code": 400, "msg": "文本不能为空"})
        result = translate_text(text)
        return jsonify({"code": 200, "data": result})
    except Exception as e:
        import traceback
        error_msg = str(e) + "\n" + traceback.format_exc()
        print(error_msg)  # 打印到日志
        return jsonify({"code": 500, "msg": "翻译失败: " + str(e)}), 500



@app.route("/api/translate/docx", methods=["POST"])
def translate_docx_api():
    """Word 文档翻译接口"""
    if 'file' not in request.files:
        return jsonify({"code": 400, "msg": "未上传文件"})
    file = request.files['file']
    if not file.filename.endswith('.docx'):
        return jsonify({"code": 400, "msg": "仅支持 .docx 格式"})
    try:
        content = read_word_file(file.read())
        if content is None:
            return jsonify({"code": 500, "msg": "python-docx 未安装，无法处理文档"})
        if not content:
            return jsonify({"code": 400, "msg": "文档内容为空"})
        result = translate_text(content)
        return jsonify({"code": 200, "data": result})
    except Exception as e:
        print(f"文档翻译错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)})

@app.route("/api/translate/export", methods=["POST"])
def export_translation():
    """导出翻译结果到 Word"""
    data = request.get_json()
    text = data.get('text', '')
    if not text:
        return jsonify({"code": 400, "msg": "文本不能为空"})
    try:
        from docx import Document
        from docx.shared import Pt
        from datetime import datetime
        
        doc = Document()
        title = doc.add_heading('翻译结果', level=1)
        title.alignment = 1  # 居中
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para)
                p.runs[0].font.size = Pt(11)
            else:
                doc.add_paragraph()
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'翻译结果_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        )
    except ImportError:
        return jsonify({"code": 500, "msg": "python-docx 未安装，无法导出"})
    except Exception as e:
        print(f"导出错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)})

@app.route("/test")
def test():
    return "服务运行正常"










# ==================== 学生管理模块 ====================
@app.route("/api/student/list")
def student_list():
    """获取学生列表"""
    try:
        db = get_db()
        cur = db.cursor()
        cur.execute("SELECT id, name, phone, grade, school FROM student ORDER BY id DESC")
        data = cur.fetchall()
        cur.close()
        db.close()
        result = [{"id": r[0], "name": r[1], "phone": r[2] or '', "grade": r[3] or '', "school": r[4] or ''} for r in
                  data]
        return jsonify({"code": 200, "data": result})
    except Exception as e:
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/student/detail/<int:student_id>", methods=["GET"])
def student_detail(student_id):
    """获取学生详情"""
    try:
        db = get_db()
        cur = db.cursor()
        cur.execute("""
            SELECT id, name, phone, grade, school, parent_name, parent_phone 
            FROM student WHERE id=%s
        """, (student_id,))
        data = cur.fetchone()
        cur.close()
        db.close()

        if data:
            return jsonify({
                "code": 200,
                "data": {
                    "id": data[0],
                    "name": data[1] or '',
                    "phone": data[2] or '',
                    "grade": data[3] or '',
                    "school": data[4] or '',
                    "parent_name": data[5] or '',
                    "parent_phone": data[6] or ''
                }
            })
        return jsonify({"code": 404, "msg": "学生不存在"})
    except Exception as e:
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/student/add", methods=["POST"])
def student_add():
    """添加学生"""
    try:
        d = request.json
        db = get_db()
        cur = db.cursor()
        cur.execute("""
            INSERT INTO student (name, phone, grade, school, parent_name, parent_phone) 
            VALUES (%s, %s, %s, %s, %s, %s)
        """, (
            d.get('name', ''),
            d.get('phone', ''),
            d.get('grade', ''),
            d.get('school', ''),
            d.get('parent_name', ''),
            d.get('parent_phone', '')
        ))
        db.commit()
        cur.close()
        db.close()
        return jsonify({"code": 200, "msg": "添加成功"})
    except Exception as e:
        return jsonify({"code": 500, "msg": f"添加失败: {str(e)}"}), 500


@app.route("/api/student/update", methods=["POST"])
def student_update():
    """更新学生信息"""
    try:
        d = request.json
        db = get_db()
        cur = db.cursor()
        cur.execute("""
            UPDATE student 
            SET name=%s, phone=%s, grade=%s, school=%s, parent_name=%s, parent_phone=%s 
            WHERE id=%s
        """, (
            d.get('name', ''),
            d.get('phone', ''),
            d.get('grade', ''),
            d.get('school', ''),
            d.get('parent_name', ''),
            d.get('parent_phone', ''),
            d.get('id')
        ))
        db.commit()
        cur.close()
        db.close()
        return jsonify({"code": 200, "msg": "更新成功"})
    except Exception as e:
        return jsonify({"code": 500, "msg": f"更新失败: {str(e)}"}), 500


@app.route("/api/student/delete", methods=["POST"])
def student_delete():
    """删除学生"""
    try:
        data = request.json
        student_id = data.get('id')
        db = get_db()
        cur = db.cursor()
        cur.execute("DELETE FROM student WHERE id=%s", (student_id,))
        db.commit()
        cur.close()
        db.close()
        return jsonify({"code": 200, "msg": "删除成功"})
    except Exception as e:
        return jsonify({"code": 500, "msg": f"删除失败: {str(e)}"}), 500










# ==================== 教师管理模块 ====================
@app.route("/api/teacher/list", methods=["GET"])
def teacher_list():
    """获取教师列表"""
    try:
        db = get_db()
        cur = db.cursor()
        cur.execute("SELECT id, name, phone, subject, class_fee FROM teacher WHERE status='active' ORDER BY id DESC")
        data = cur.fetchall()
        cur.close()
        db.close()
        result = [{"id": r[0], "name": r[1] or '', "phone": r[2] or '', "subject": r[3] or '',
                   "class_fee": float(r[4]) if r[4] else 0} for r in data]
        return jsonify({"code": 200, "data": result})
    except Exception as e:
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/teacher/detail/<int:teacher_id>", methods=["GET"])
def teacher_detail(teacher_id):
    """获取教师详情"""
    try:
        db = get_db()
        cur = db.cursor()
        cur.execute("SELECT id, name, phone, subject, class_fee FROM teacher WHERE id=%s", (teacher_id,))
        data = cur.fetchone()
        cur.close()
        db.close()

        if data:
            return jsonify({
                "code": 200,
                "data": {
                    "id": data[0],
                    "name": data[1] or '',
                    "phone": data[2] or '',
                    "subject": data[3] or '',
                    "class_fee": float(data[4]) if data[4] else 0
                }
            })
        return jsonify({"code": 404, "msg": "教师不存在"})
    except Exception as e:
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/teacher/add", methods=["POST"])
def teacher_add():
    """添加教师"""
    try:
        d = request.json
        db = get_db()
        cur = db.cursor()
        cur.execute("""
            INSERT INTO teacher (name, phone, subject, class_fee, status) 
            VALUES (%s, %s, %s, %s, 'active')
        """, (d.get('name', ''), d.get('phone', ''), d.get('subject', ''), d.get('class_fee', 0)))
        db.commit()
        cur.close()
        db.close()
        return jsonify({"code": 200, "msg": "添加成功"})
    except Exception as e:
        return jsonify({"code": 500, "msg": f"添加失败: {str(e)}"}), 500


@app.route("/api/teacher/update", methods=["POST"])
def teacher_update():
    """更新教师信息"""
    try:
        d = request.json
        db = get_db()
        cur = db.cursor()
        cur.execute("""
            UPDATE teacher 
            SET name=%s, phone=%s, subject=%s, class_fee=%s 
            WHERE id=%s
        """, (d.get('name', ''), d.get('phone', ''), d.get('subject', ''), d.get('class_fee', 0), d.get('id')))
        db.commit()
        cur.close()
        db.close()
        return jsonify({"code": 200, "msg": "更新成功"})
    except Exception as e:
        return jsonify({"code": 500, "msg": f"更新失败: {str(e)}"}), 500


@app.route("/api/teacher/delete", methods=["POST"])
def teacher_delete():
    """删除教师（软删除）"""
    try:
        data = request.json
        teacher_id = data.get('id')
        db = get_db()
        cur = db.cursor()
        cur.execute("UPDATE teacher SET status='inactive' WHERE id=%s", (teacher_id,))
        db.commit()
        cur.close()
        db.close()
        return jsonify({"code": 200, "msg": "删除成功"})
    except Exception as e:
        return jsonify({"code": 500, "msg": f"删除失败: {str(e)}"}), 500


@app.route("/api/teacher/salary/statistics", methods=["GET"])
def teacher_salary_statistics():
    """教师课酬统计"""
    try:
        teacher_id = request.args.get('teacher_id', type=int)
        month = request.args.get('month')

        db = get_db()
        cur = db.cursor()

        sql = """
            SELECT 
                t.id as teacher_id,
                t.name as teacher_name,
                t.subject,
                t.class_fee,
                COUNT(CASE WHEN s.status = 'completed' THEN 1 END) as completed_classes,
                COALESCE(SUM(CASE WHEN s.status = 'completed' THEN s.duration ELSE 0 END), 0) as total_hours,
                COALESCE(SUM(CASE WHEN s.status = 'completed' THEN s.duration * t.class_fee ELSE 0 END), 0) as total_amount
            FROM teacher t
            LEFT JOIN course_schedule s ON t.id = s.teacher_id
            WHERE t.status = 'active'
        """
        params = []

        if teacher_id:
            sql += " AND t.id = %s"
            params.append(teacher_id)
        if month:
            sql += " AND TO_CHAR(s.class_date, 'YYYY-MM') = %s"
            params.append(month)

        sql += " GROUP BY t.id, t.name, t.subject, t.class_fee ORDER BY t.id"

        cur.execute(sql, params)
        data = cur.fetchall()
        cur.close()
        db.close()

        result = []
        for r in data:
            result.append({
                "teacher_id": r[0],
                "teacher_name": r[1] or '',
                "subject": r[2] or '',
                "class_fee": float(r[3]) if r[3] else 0,
                "completed_classes": r[4] or 0,
                "total_hours": float(r[5]) if r[5] else 0,
                "total_amount": float(r[6]) if r[6] else 0
            })

        return jsonify({"code": 200, "data": result})
    except Exception as e:
        return jsonify({"code": 500, "msg": str(e)}), 500



# ==================== 课时管理模块 ====================
@app.route("/api/student/hours/list", methods=["GET"])
def student_hours_list():
    """获取学生课时列表"""
    try:
        student_id = request.args.get('student_id', type=int)
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        db = get_db()
        cur = db.cursor()

        sql = """
            SELECT 
                cp.id,
                s.id as student_id,
                s.name as student_name,
                s.grade,
                s.phone,
                cp.course_name,
                cp.total,
                cp.used,
                cp.surplus,
                cp.expire_date,
                cp.status,
                cp.created_at
            FROM course_package cp
            JOIN student s ON cp.student_id = s.id
            WHERE cp.status = 'active'
        """
        params = []

        if student_id:
            sql += " AND cp.student_id = %s"
            params.append(student_id)
        if start_date:
            sql += " AND cp.created_at >= %s"
            params.append(start_date)
        if end_date:
            sql += " AND cp.created_at <= %s"
            params.append(end_date)

        sql += " ORDER BY cp.created_at DESC"

        cur.execute(sql, params)
        data = cur.fetchall()
        cur.close()
        db.close()

        result = []
        for r in data:
            result.append({
                "id": r[0],
                "student_id": r[1],
                "student_name": r[2],
                "grade": r[3] or '',
                "phone": r[4] or '',
                "course_name": r[5] or '标准课程',
                "total": float(r[6]) if r[6] else 0,
                "used": float(r[7]) if r[7] else 0,
                "surplus": float(r[8]) if r[8] else 0,
                "expire_date": str(r[9]) if r[9] else None,
                "status": r[10] or 'active',
                "created_at": str(r[11]) if r[11] else None
            })

        return jsonify({"code": 200, "data": result})
    except Exception as e:
        print(f"获取课时列表错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/student/hours/statistics", methods=["GET"])
def student_hours_statistics():
    """获取课时统计数据"""
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        db = get_db()
        cur = db.cursor()

        # 总课时统计
        sql = """
            SELECT 
                COUNT(*) as total_packages,
                COALESCE(SUM(total), 0) as total_hours,
                COALESCE(SUM(used), 0) as used_hours,
                COALESCE(SUM(surplus), 0) as surplus_hours
            FROM course_package
            WHERE status = 'active'
        """
        params = []
        if start_date:
            sql += " AND created_at >= %s"
            params.append(start_date)
        if end_date:
            sql += " AND created_at <= %s"
            params.append(end_date)

        cur.execute(sql, params)
        total_stats = cur.fetchone()

        # 即将过期课时
        cur.execute("""
            SELECT 
                COUNT(*) as expiring_count,
                COALESCE(SUM(surplus), 0) as expiring_hours
            FROM course_package
            WHERE status = 'active'
              AND expire_date IS NOT NULL
              AND expire_date <= CURRENT_DATE + INTERVAL '30 days'
        """)
        expiring_stats = cur.fetchone()

        cur.close()
        db.close()

        return jsonify({"code": 200, "data": {
            "total_packages": total_stats[0] or 0,
            "total_hours": float(total_stats[1]) if total_stats[1] else 0,
            "used_hours": float(total_stats[2]) if total_stats[2] else 0,
            "surplus_hours": float(total_stats[3]) if total_stats[3] else 0,
            "expiring_count": expiring_stats[0] or 0,
            "expiring_hours": float(expiring_stats[1]) if expiring_stats[1] else 0
        }})
    except Exception as e:
        print(f"获取课时统计错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/student/hours/statistics-v2", methods=["GET"])
def get_student_hours_statistics_v2():
    """获取学生课时消耗统计（基于已确认课程）"""
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        db = get_db()
        cur = db.cursor()
        
        # 基于 course_schedule 表统计已确认课程
        sql = """
            SELECT 
                s.id as student_id,
                s.name as student_name,
                s.grade,
                COUNT(cs.id) as class_count,
                COALESCE(SUM(cs.duration), 0) as total_hours
            FROM student s
            LEFT JOIN course_schedule cs ON 
                (cs.student_id = s.id 
                 OR cs.student_ids = CAST(s.id AS TEXT)
                 OR cs.student_ids LIKE CONCAT(CAST(s.id AS TEXT), ',%')
                 OR cs.student_ids LIKE CONCAT('%,', CAST(s.id AS TEXT), ',%'))
                AND cs.status = 'completed'
        """
        params = []
        
        if start_date and end_date:
            sql += " AND cs.class_date BETWEEN %s AND %s"
            params.extend([start_date, end_date])
        
        sql += """
            GROUP BY s.id, s.name, s.grade
            HAVING COUNT(cs.id) > 0
            ORDER BY s.name
        """
        
        cur.execute(sql, params)
        
        student_stats = []
        total_classes = 0
        total_hours = 0
        
        for row in cur.fetchall():
            student_stats.append({
                "student_id": row[0],
                "student_name": row[1] or '',
                "grade": row[2] or '',
                "class_count": row[3] or 0,
                "total_hours": float(row[4]) if row[4] else 0
            })
            total_classes += row[3] or 0
            total_hours += float(row[4]) if row[4] else 0
        
        cur.close()
        db.close()
        
        return jsonify({
            "code": 200,
            "data": {
                "total_students": len(student_stats),
                "total_classes": total_classes,
                "total_hours": total_hours,
                "student_stats": student_stats
            }
        })
    except Exception as e:
        print(f"获取统计错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500



@app.route("/api/student/hours/add", methods=["POST"])
def student_hours_add():
    """添加课时包"""
    try:
        data = request.json
        student_id = data.get('student_id')
        total_hours = data.get('total_hours')
        course_name = data.get('course_name', '标准课程')
        expire_date = data.get('expire_date')

        if not student_id:
            return jsonify({"code": 400, "msg": "请选择学生"}), 400
        if not total_hours or total_hours <= 0:
            return jsonify({"code": 400, "msg": "请输入有效的课时数"}), 400

        db = get_db()
        cur = db.cursor()

        cur.execute("""
            INSERT INTO course_package (student_id, total, used, surplus, course_name, expire_date, status)
            VALUES (%s, %s, 0, %s, %s, %s, 'active')
            RETURNING id
        """, (student_id, total_hours, total_hours, course_name, expire_date))

        new_id = cur.fetchone()[0]
        db.commit()
        cur.close()
        db.close()

        return jsonify({"code": 200, "msg": "添加成功", "data": {"id": new_id}})
    except Exception as e:
        print(f"添加课时包错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/student/hours/consume", methods=["POST"])
def student_hours_consume():
    """消耗课时"""
    try:
        data = request.json
        package_id = data.get('package_id')
        hours = data.get('hours')
        note = data.get('note', '')

        if not package_id:
            return jsonify({"code": 400, "msg": "请选择课时包"}), 400
        if not hours or hours <= 0:
            return jsonify({"code": 400, "msg": "请输入有效的消耗课时"}), 400

        db = get_db()
        cur = db.cursor()

        # 检查剩余课时
        cur.execute("SELECT surplus FROM course_package WHERE id = %s", (package_id,))
        result = cur.fetchone()
        if not result:
            return jsonify({"code": 404, "msg": "课时包不存在"}), 404

        surplus = result[0]
        if surplus < hours:
            return jsonify({"code": 400, "msg": f"剩余课时不足，仅剩 {surplus} 小时"}), 400

        # 更新课时包
        cur.execute("""
            UPDATE course_package 
            SET used = used + %s, surplus = surplus - %s
            WHERE id = %s
        """, (hours, hours, package_id))

        # 记录消耗日志
        cur.execute("""
            INSERT INTO hour_consumption (package_id, hours, consume_date, note)
            VALUES (%s, %s, CURRENT_DATE, %s)
        """, (package_id, hours, note))

        db.commit()
        cur.close()
        db.close()

        return jsonify({"code": 200, "msg": f"成功消耗 {hours} 课时"})
    except Exception as e:
        print(f"消耗课时错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/student/hours/consumption/list", methods=["GET"])
def hours_consumption_list():
    """获取课时消耗记录"""
    try:
        package_id = request.args.get('package_id', type=int)

        db = get_db()
        cur = db.cursor()

        sql = """
            SELECT 
                hc.id,
                hc.package_id,
                cp.course_name,
                s.name as student_name,
                hc.hours,
                hc.consume_date,
                hc.note,
                hc.created_at
            FROM hour_consumption hc
            LEFT JOIN course_package cp ON hc.package_id = cp.id
            LEFT JOIN student s ON cp.student_id = s.id
            WHERE 1=1
        """
        params = []

        if package_id:
            sql += " AND hc.package_id = %s"
            params.append(package_id)

        sql += " ORDER BY hc.consume_date DESC, hc.created_at DESC LIMIT 100"

        cur.execute(sql, params)
        data = cur.fetchall()
        cur.close()
        db.close()

        result = []
        for r in data:
            result.append({
                "id": r[0],
                "package_id": r[1],
                "course_name": r[2] or '',
                "student_name": r[3] or '',
                "hours": float(r[4]) if r[4] else 0,
                "consume_date": str(r[5]) if r[5] else None,
                "note": r[6] or ''
            })

        return jsonify({"code": 200, "data": result})
    except Exception as e:
        print(f"获取消耗记录错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/student/hours/export", methods=["POST"])
def student_hours_export():
    """导出课时报表"""
    try:
        data = request.json
        student_ids = data.get('student_ids', [])

        db = get_db()
        cur = db.cursor()

        if student_ids:
            placeholders = ','.join(['%s'] * len(student_ids))
            sql = f"""
                SELECT 
                    s.name,
                    s.grade,
                    s.phone,
                    cp.course_name,
                    cp.total,
                    cp.used,
                    cp.surplus,
                    cp.expire_date
                FROM course_package cp
                JOIN student s ON cp.student_id = s.id
                WHERE cp.status = 'active' AND s.id IN ({placeholders})
                ORDER BY s.grade, s.name
            """
            cur.execute(sql, student_ids)
        else:
            cur.execute("""
                SELECT 
                    s.name,
                    s.grade,
                    s.phone,
                    cp.course_name,
                    cp.total,
                    cp.used,
                    cp.surplus,
                    cp.expire_date
                FROM course_package cp
                JOIN student s ON cp.student_id = s.id
                WHERE cp.status = 'active'
                ORDER BY s.grade, s.name
            """)

        data = cur.fetchall()
        cur.close()
        db.close()

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "学生课时统计"

        headers = ["学生姓名", "年级", "联系电话", "课程名称", "总课时", "已用课时", "剩余课时", "有效期"]
        ws.append(headers)

        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        for row in data:
            ws.append([
                row[0] or '',
                row[1] or '',
                row[2] or '',
                row[3] or '',
                float(row[4]) if row[4] else 0,
                float(row[5]) if row[5] else 0,
                float(row[6]) if row[6] else 0,
                str(row[7]) if row[7] else '永久'
            ])

        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 25)
            ws.column_dimensions[col_letter].width = adjusted_width

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'学生课时统计_{datetime.now().strftime("%Y%m%d")}.xlsx'
        )
    except Exception as e:
        print(f"导出课时报表错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


# ==================== 学期统计模块 ====================

@app.route("/api/student/semester/list", methods=["GET"])
def get_semester_list():
    """获取学期列表（基于排课数据）"""
    try:
        db = get_db()
        cur = db.cursor()
        
        cur.execute("""
            SELECT DISTINCT 
                TO_CHAR(class_date, 'YYYY') as year,
                CASE 
                    WHEN EXTRACT(MONTH FROM class_date) BETWEEN 3 AND 8 THEN 'spring'
                    ELSE 'autumn'
                END as semester_type,
                CASE 
                    WHEN EXTRACT(MONTH FROM class_date) BETWEEN 3 AND 8 THEN '春季学期'
                    ELSE '秋季学期'
                END as semester_name
            FROM course_schedule 
            WHERE status = 'completed'
            ORDER BY year DESC, semester_type DESC
        """)
        
        data = cur.fetchall()
        cur.close()
        db.close()
        
        result = []
        for row in data:
            result.append({
                "key": f"{row[0]}-{row[1]}",
                "name": f"{row[0]}{row[2]}",
                "year": row[0],
                "semester_type": row[1]
            })
        
        return jsonify({"code": 200, "data": result})
    except Exception as e:
        print(f"获取学期列表错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/student/semester/statistics", methods=["GET"])
def get_semester_statistics():
    """获取学期统计数据（支持多学生）"""
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        if not start_date or not end_date:
            return jsonify({"code": 400, "msg": "缺少时间范围参数"}), 400
        
        db = get_db()
        cur = db.cursor()
        
        # 修复：正确统计每个学生的上课次数（支持多学生）
        # 使用子查询先找出每个学生关联的课程，再统计
        cur.execute("""
            WITH student_courses AS (
                SELECT 
                    s.id as student_id,
                    s.name as student_name,
                    s.grade,
                    cs.id as course_id,
                    cs.duration
                FROM student s
                LEFT JOIN course_schedule cs ON 
                    cs.status = 'completed'
                    AND cs.class_date BETWEEN %s AND %s
                    AND (
                        cs.student_id = s.id 
                        OR cs.student_ids = CAST(s.id AS TEXT)
                        OR cs.student_ids LIKE CONCAT(CAST(s.id AS TEXT), ',%')
                        OR cs.student_ids LIKE CONCAT('%,', CAST(s.id AS TEXT), ',%')
                        OR cs.student_ids LIKE CONCAT('%,', CAST(s.id AS TEXT))
                    )
            )
            SELECT 
                student_id,
                student_name,
                grade,
                COUNT(DISTINCT course_id) as class_count,
                COALESCE(SUM(duration), 0) as total_hours
            FROM student_courses
            WHERE course_id IS NOT NULL
            GROUP BY student_id, student_name, grade
            ORDER BY student_name
        """, (start_date, end_date))
        
        student_stats = []
        total_classes = 0
        total_hours = 0
        
        for row in cur.fetchall():
            student_stats.append({
                "student_id": row[0],
                "student_name": row[1] or '',
                "grade": row[2] or '',
                "class_count": row[3] or 0,
                "total_hours": float(row[4]) if row[4] else 0
            })
            total_classes += row[3] or 0
            total_hours += float(row[4]) if row[4] else 0
        
        cur.close()
        db.close()
        
        return jsonify({
            "code": 200,
            "data": {
                "start_date": start_date,
                "end_date": end_date,
                "total_students": len(student_stats),
                "total_classes": total_classes,
                "total_hours": total_hours,
                "student_stats": student_stats
            }
        })
    except Exception as e:
        print(f"获取学期统计错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/student/semester/detail", methods=["GET"])
def get_student_semester_detail():
    """获取学生在指定时间范围内的课程明细（支持多学生）"""
    try:
        student_id = request.args.get('student_id', type=int)
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        if not student_id or not start_date or not end_date:
            return jsonify({"code": 400, "msg": "缺少参数"}), 400
        
        db = get_db()
        cur = db.cursor()
        
        student_id_str = str(student_id)
        
        # 修复：正确匹配单学生和多学生情况
        cur.execute("""
            SELECT 
                cs.id,
                cs.class_date,
                cs.class_time,
                cs.subject,
                cs.classroom,
                cs.duration,
                t.name as teacher_name
            FROM course_schedule cs
            LEFT JOIN teacher t ON cs.teacher_id = t.id
            WHERE cs.status = 'completed'
              AND cs.class_date BETWEEN %s AND %s
              AND (
                  cs.student_id = %s 
                  OR cs.student_ids = %s
                  OR cs.student_ids LIKE %s
                  OR cs.student_ids LIKE %s
                  OR cs.student_ids LIKE %s
              )
            ORDER BY cs.class_date, cs.class_time
        """, (
            start_date, end_date, 
            student_id, 
            student_id_str,
            f'{student_id_str},%',
            f'%,{student_id_str},%',
            f'%,{student_id_str}'
        ))
        
        schedule_list = []
        total_classes = 0
        total_hours = 0
        
        for row in cur.fetchall():
            duration = float(row[5]) if row[5] else 2
            schedule_list.append({
                "id": row[0],
                "class_date": str(row[1]),
                "class_time": row[2],
                "subject": row[3] or '',
                "classroom": row[4] or '',
                "duration": duration,
                "teacher_name": row[6] or '待分配'
            })
            total_classes += 1
            total_hours += duration
        
        cur.execute("SELECT name FROM student WHERE id = %s", (student_id,))
        student = cur.fetchone()
        student_name = student[0] if student else ''
        
        cur.close()
        db.close()
        
        return jsonify({
            "code": 200,
            "data": {
                "student_id": student_id,
                "student_name": student_name,
                "total_classes": total_classes,
                "total_hours": total_hours,
                "schedule_list": schedule_list
            }
        })
    except Exception as e:
        print(f"获取学生明细错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500



@app.route("/api/student/semester/export", methods=["POST"])
def export_semester_report():
    """导出学期统计报表"""
    try:
        data = request.json
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        student_ids = data.get('student_ids', [])
        
        if not start_date or not end_date:
            return jsonify({"code": 400, "msg": "缺少时间范围参数"}), 400
        
        db = get_db()
        cur = db.cursor()
        
        # 获取统计数据
        if student_ids:
            placeholders = ','.join(['%s'] * len(student_ids))
            sql = f"""
                SELECT 
                    s.id,
                    s.name,
                    s.grade,
                    COUNT(cs.id) as class_count,
                    COALESCE(SUM(cs.duration), 0) as total_hours
                FROM student s
                LEFT JOIN course_schedule cs ON 
                    (cs.student_id = s.id 
                     OR cs.student_ids = CAST(s.id AS TEXT)
                     OR cs.student_ids LIKE CONCAT(CAST(s.id AS TEXT), ',%')
                     OR cs.student_ids LIKE CONCAT('%,', CAST(s.id AS TEXT), ',%'))
                    AND cs.status = 'completed'
                    AND cs.class_date BETWEEN %s AND %s
                WHERE s.id IN ({placeholders})
                GROUP BY s.id, s.name, s.grade
                ORDER BY s.name
            """
            params = [start_date, end_date] + student_ids
            cur.execute(sql, params)
        else:
            cur.execute("""
                SELECT 
                    s.id,
                    s.name,
                    s.grade,
                    COUNT(cs.id) as class_count,
                    COALESCE(SUM(cs.duration), 0) as total_hours
                FROM student s
                LEFT JOIN course_schedule cs ON 
                    (cs.student_id = s.id 
                     OR cs.student_ids = CAST(s.id AS TEXT)
                     OR cs.student_ids LIKE CONCAT(CAST(s.id AS TEXT), ',%')
                     OR cs.student_ids LIKE CONCAT('%,', CAST(s.id AS TEXT), ',%'))
                    AND cs.status = 'completed'
                    AND cs.class_date BETWEEN %s AND %s
                GROUP BY s.id, s.name, s.grade
                HAVING COUNT(cs.id) > 0
                ORDER BY s.name
            """, (start_date, end_date))
        
        data = cur.fetchall()
        cur.close()
        db.close()
        
        # 创建 Excel 文件
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = f"课时统计_{start_date}_至_{end_date}"
        
        headers = ["学生姓名", "年级", "上课次数", "总课时(小时)"]
        ws.append(headers)
        
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")
        
        total_classes = 0
        total_hours = 0
        
        for row in data:
            ws.append([
                row[1] or '',
                row[2] or '',
                row[3] or 0,
                float(row[4]) if row[4] else 0
            ])
            total_classes += row[3] or 0
            total_hours += float(row[4]) if row[4] else 0
        
        # 添加汇总行
        ws.append([])
        ws.append(["合计", "", total_classes, total_hours])
        
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 25)
            ws.column_dimensions[col_letter].width = adjusted_width
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'课时统计_{start_date}_至_{end_date}.xlsx'
        )
    except Exception as e:
        print(f"导出报表错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500




# ==================== 排课管理模块 ====================
@app.route("/api/schedule/calendar", methods=["GET"])
def get_schedule_calendar():
    try:
        date_str = request.args.get('date')
        if not date_str:
            return jsonify({"code": 400, "msg": "缺少日期参数"}), 400

        db = get_db()
        cur = db.cursor()
        # 查询课程基本信息
        cur.execute("""
            SELECT 
                s.id,
                s.class_time,
                s.subject,
                s.classroom,
                COALESCE(s.status, 'scheduled') as status,
                t.name as teacher_name,
                s.student_ids
            FROM course_schedule s
            LEFT JOIN teacher t ON s.teacher_id = t.id
            WHERE s.class_date = %s 
              AND (s.status IS NULL OR s.status != 'cancelled')
            ORDER BY s.class_time
        """, (date_str,))
        
        rows = cur.fetchall()
        # 获取学生名称映射
        cur.execute("SELECT id, name FROM student")
        student_map = {row[0]: row[1] for row in cur.fetchall()}
        cur.close()
        db.close()

        schedule_list = []
        for row in rows:
            student_ids_str = row[6] or ''
            student_ids = [int(x) for x in student_ids_str.split(',') if x]
            student_names = [student_map.get(sid, '') for sid in student_ids if student_map.get(sid)]
            schedule_list.append({
                "id": row[0],
                "class_time": row[1],
                "subject": row[2],
                "classroom": row[3] or '',
                "status": row[4],
                "teacher_name": row[5] or '待分配',
                "students": student_names,          # 新增学生姓名数组
                "student_count": len(student_names) # 也可单独返回人数
            })
        return jsonify({"code": 200, "data": schedule_list})
    except Exception as e:
        print(f"获取日历数据错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/schedule/week", methods=["GET"])
def get_week_schedule():
    """获取周课表数据 - 每个时间段支持多课程"""
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        if not start_date or not end_date:
            return jsonify({"code": 400, "msg": "缺少日期参数"}), 400
        
        db = get_db()
        cur = db.cursor()
        
        # 查询所有课程
        cur.execute("""
            SELECT 
                s.id,
                s.class_date,
                EXTRACT(DOW FROM s.class_date) as weekday,
                s.class_time,
                s.subject,
                s.classroom,
                COALESCE(s.status, 'scheduled') as status,
                COALESCE(s.duration, 2) as duration,
                t.name as teacher_name,
                s.student_ids
            FROM course_schedule s
            LEFT JOIN teacher t ON s.teacher_id = t.id
            WHERE s.class_date BETWEEN %s AND %s
              AND (s.status IS NULL OR s.status != 'cancelled')
            ORDER BY s.class_date, s.class_time, s.id
        """, (start_date, end_date))

        rows = cur.fetchall()

        # 修正：获取学生名称映射（字段名应为 'name'，而非 'student_name'）
        cur.execute("SELECT id, name FROM student")
        # 如果表中字段是 'student_name'，请改为 'student_name'
        students_map = {row[0]: row[1] for row in cur.fetchall()}

        cur.close()
        db.close()

        week_schedule = {}
        for row in rows:
            weekday = int(row[2])
            weekday_idx = 6 if weekday == 0 else weekday - 1
            time_slot = row[3]

            if weekday_idx not in week_schedule:
                week_schedule[weekday_idx] = {}
            if time_slot not in week_schedule[weekday_idx]:
                week_schedule[weekday_idx][time_slot] = []

            # 解析 student_ids 并获取姓名
            student_ids_str = row[9] or ''
            student_id_list = [int(x) for x in student_ids_str.split(',') if x]
            student_names = [students_map.get(sid, '') for sid in student_id_list if students_map.get(sid)]

            week_schedule[weekday_idx][time_slot].append({
                "id": row[0],
                "subject": row[4] or '',
                "teacher": row[8] or '',
                "place": row[5] or '',
                "duration": float(row[7]) if row[7] else 2,
                "status": row[6],
                "students": student_names
            })

        return jsonify({"code": 200, "data": week_schedule})
    except Exception as e:
        print(f"周课表错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500





@app.route("/api/schedule/detail/<int:schedule_id>", methods=["GET"])
def schedule_detail(schedule_id):
    """获取排课详情"""
    try:
        db = get_db()
        cur = db.cursor()

        cur.execute("""
            SELECT 
                s.id,
                s.student_id,
                s.teacher_id,
                s.class_date,
                s.class_time,
                s.subject,
                s.classroom,
                s.status,
                s.duration,
                s.student_ids,
                t.name as teacher_name
            FROM course_schedule s
            LEFT JOIN teacher t ON s.teacher_id = t.id
            WHERE s.id = %s
        """, (schedule_id,))

        data = cur.fetchone()
        cur.close()
        db.close()

        if data:
            # 解析学生ID列表
            student_id_list = []
            if data[9]:
                student_id_list = [int(x) for x in data[9].split(',') if x]

            return jsonify({
                "code": 200,
                "data": {
                    "id": data[0],
                    "student_id": data[1],
                    "teacher_id": data[2],
                    "class_date": str(data[3]),
                    "class_time": data[4],
                    "subject": data[5],
                    "classroom": data[6] or '',
                    "status": data[7] or 'scheduled',
                    "duration": float(data[8]) if data[8] else 2,
                    "student_ids": data[9] or '',
                    "student_id_list": student_id_list,
                    "teacher_name": data[10] or ''
                }
            })
        return jsonify({"code": 404, "msg": "排课不存在"})
    except Exception as e:
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/schedule/save", methods=["POST"])
def schedule_save():
    """保存排课 - 支持同一时间段多课程"""
    try:
        data = request.json
        print("=== 收到保存请求 ===")
        print(json.dumps(data, ensure_ascii=False))
        
        subject = data.get('subject')
        teacher_id = data.get('teacher_id')
        student_ids = data.get('student_ids', '')
        classroom = data.get('classroom')
        class_date = data.get('date')
        class_time = data.get('time')
        duration = data.get('duration', 2)
        
        if not subject:
            return jsonify({"code": 400, "msg": "课程名称不能为空"}), 400
        if not class_date:
            return jsonify({"code": 400, "msg": "日期不能为空"}), 400
        if not class_time:
            return jsonify({"code": 400, "msg": "时间不能为空"}), 400
        
        if teacher_id:
            try:
                teacher_id = int(teacher_id)
            except (ValueError, TypeError):
                teacher_id = None
        
        db = get_db()
        cur = db.cursor()
        
        # 重要：不再检查是否已存在，直接插入新课程
        # 这样同一时间段可以有多门课程
        cur.execute("""
            INSERT INTO course_schedule 
            (subject, teacher_id, student_ids, classroom, class_date, class_time, duration, status)
            VALUES (%s, %s, %s, %s, %s, %s, %s, 'scheduled')
            RETURNING id
        """, (subject, teacher_id, student_ids, classroom, class_date, class_time, duration))
        
        new_id = cur.fetchone()[0]
        db.commit()
        
        print(f"添加成功，新课程ID: {new_id}")
        
        cur.close()
        db.close()
        
        return jsonify({"code": 200, "msg": "添加成功", "data": {"id": new_id}})
    except Exception as e:
        print(f"保存排课错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500

@app.route("/api/schedule/update", methods=["POST"])
def schedule_update():
    """更新排课信息 - 只更新指定ID的课程"""
    try:
        data = request.json
        print("=== 收到更新请求 ===")
        print(json.dumps(data, ensure_ascii=False))
        
        schedule_id = data.get('id')
        subject = data.get('subject')
        teacher_id = data.get('teacher_id')
        student_ids = data.get('student_ids', '')
        classroom = data.get('classroom')
        class_date = data.get('date')
        class_time = data.get('time')
        duration = data.get('duration', 2)
        
        if not schedule_id:
            return jsonify({"code": 400, "msg": "缺少课程ID"}), 400
        
        if teacher_id:
            try:
                teacher_id = int(teacher_id)
            except (ValueError, TypeError):
                teacher_id = None
        
        db = get_db()
        cur = db.cursor()
        
        # 只更新指定ID的课程，不影响同一时间段的其他课程
        cur.execute("""
            UPDATE course_schedule 
            SET subject = %s, teacher_id = %s, student_ids = %s, 
                classroom = %s, class_date = %s, class_time = %s, duration = %s
            WHERE id = %s
        """, (subject, teacher_id, student_ids, classroom, class_date, class_time, duration, schedule_id))
        
        db.commit()
        print(f"更新成功，ID: {schedule_id}")
        
        cur.close()
        db.close()
        
        return jsonify({"code": 200, "msg": "更新成功"})
    except Exception as e:
        print(f"更新排课错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/schedule/delete", methods=["POST"])
def schedule_delete():
    """删除排课（软删除）"""
    try:
        data = request.json
        schedule_id = data.get('id')

        db = get_db()
        cur = db.cursor()

        cur.execute("UPDATE course_schedule SET status = 'cancelled' WHERE id = %s", (schedule_id,))

        db.commit()
        cur.close()
        db.close()

        return jsonify({"code": 200, "msg": "取消成功"})
    except Exception as e:
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/schedule/clear_week", methods=["POST"])
def clear_week_schedule():
    """清空一周的课程"""
    try:
        data = request.json
        start_date = data.get('start_date')
        end_date = data.get('end_date')

        db = get_db()
        cur = db.cursor()

        cur.execute("""
            UPDATE course_schedule 
            SET status = 'cancelled' 
            WHERE class_date BETWEEN %s AND %s
        """, (start_date, end_date))

        db.commit()
        cur.close()
        db.close()

        return jsonify({"code": 200, "msg": "清空成功"})
    except Exception as e:
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/schedule/copy_week", methods=["POST"])
def copy_week_schedule():
    """复制本周课程到下一周"""
    try:
        data = request.json
        current_start_date = data.get('current_start_date')
        current_end_date = data.get('current_end_date')

        if not current_start_date or not current_end_date:
            return jsonify({"code": 400, "msg": "缺少日期参数"}), 400

        db = get_db()
        cur = db.cursor()

        cur.execute("""
            SELECT 
                id,
                student_id,
                teacher_id,
                subject,
                classroom,
                class_date,
                class_time,
                duration,
                student_ids
            FROM course_schedule
            WHERE class_date BETWEEN %s AND %s
              AND (status IS NULL OR status != 'cancelled')
        """, (current_start_date, current_end_date))

        courses = cur.fetchall()

        if not courses:
            cur.close()
            db.close()
            return jsonify({"code": 200, "msg": "本周没有课程可复制", "count": 0})

        print(f"找到 {len(courses)} 门课程待复制")

        from datetime import datetime, timedelta

        copied_count = 0
        skipped_count = 0

        for course in courses:
            student_id = course[1]
            teacher_id = course[2]
            subject = course[3]
            classroom = course[4]
            old_date = course[5]
            class_time = course[6]
            duration = course[7]
            student_ids = course[8]

            new_date = old_date + timedelta(days=7)
            new_date_str = new_date.strftime("%Y-%m-%d")

            # 检查下一周是否已有相同时间段的课程
            cur.execute("""
                SELECT id FROM course_schedule
                WHERE class_date = %s AND class_time = %s
                AND (status IS NULL OR status != 'cancelled')
            """, (new_date_str, class_time))

            existing = cur.fetchone()

            if existing:
                skipped_count += 1
                continue

            cur.execute("""
                INSERT INTO course_schedule 
                (student_id, teacher_id, subject, classroom, class_date, class_time, duration, student_ids, status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, 'scheduled')
                RETURNING id
            """, (student_id, teacher_id, subject, classroom, new_date_str, class_time, duration, student_ids))

            copied_count += 1

        db.commit()
        cur.close()
        db.close()

        if copied_count == 0:
            msg = "没有课程可复制到下一周"
        else:
            msg = f"成功复制 {copied_count} 门课程到下一周"
            if skipped_count > 0:
                msg += f"，跳过 {skipped_count} 门（下一周已有相同时间段课程）"

        return jsonify({
            "code": 200,
            "msg": msg,
            "count": copied_count,
            "skipped": skipped_count
        })
    except Exception as e:
        print(f"复制周课表错误: {str(e)}")
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500


# ==================== 复制课程接口 ====================

@app.route("/api/copy/single", methods=["POST"])
def copy_single_course():
    """复制单个课程到目标日期"""
    try:
        data = request.json
        course_id = data.get('course_id')
        target_date = data.get('target_date')
        
        if not course_id or not target_date:
            return jsonify({"code": 400, "msg": "缺少参数"}), 400
        
        db = get_db()
        cur = db.cursor()
        
        # 获取原课程信息
        cur.execute("""
            SELECT subject, teacher_id, student_ids, classroom, class_time, duration
            FROM course_schedule WHERE id = %s
        """, (course_id,))
        course = cur.fetchone()
        
        if not course:
            return jsonify({"code": 404, "msg": "原课程不存在"}), 404
        
        # 插入新课程
        cur.execute("""
            INSERT INTO course_schedule 
            (subject, teacher_id, student_ids, classroom, class_date, class_time, duration, status)
            VALUES (%s, %s, %s, %s, %s, %s, %s, 'scheduled')
            RETURNING id
        """, (course[0], course[1], course[2], course[3], target_date, course[4], course[5]))
        
        new_id = cur.fetchone()[0]
        db.commit()
        cur.close()
        db.close()
        
        return jsonify({"code": 200, "msg": "复制成功", "count": 1, "id": new_id})
    except Exception as e:
        print(f"复制单个课程错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/copy/day", methods=["POST"])
def copy_day_courses():
    """复制某天的全部课程到目标日期"""
    try:
        data = request.json
        source_date = data.get('source_date')
        target_date = data.get('target_date')
        
        if not source_date or not target_date:
            return jsonify({"code": 400, "msg": "缺少参数"}), 400
        
        db = get_db()
        cur = db.cursor()
        
        # 获取源日期的所有课程
        cur.execute("""
            SELECT subject, teacher_id, student_ids, classroom, class_time, duration
            FROM course_schedule
            WHERE class_date = %s
              AND (status IS NULL OR status != 'cancelled')
        """, (source_date,))
        
        courses = cur.fetchall()
        copied_count = 0
        
        for course in courses:
            cur.execute("""
                INSERT INTO course_schedule 
                (subject, teacher_id, student_ids, classroom, class_date, class_time, duration, status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, 'scheduled')
            """, (course[0], course[1], course[2], course[3], target_date, course[4], course[5]))
            copied_count += 1
        
        db.commit()
        cur.close()
        db.close()
        
        return jsonify({"code": 200, "msg": "复制成功", "count": copied_count})
    except Exception as e:
        print(f"复制日课程错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/copy/week", methods=["POST"])
def copy_week_courses():
    """复制本周全部课程到目标日期（按星期偏移）"""
    try:
        data = request.json
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        target_date = data.get('target_date')
        
        if not start_date or not end_date or not target_date:
            return jsonify({"code": 400, "msg": "缺少参数"}), 400
        
        from datetime import datetime, timedelta
        target = datetime.strptime(target_date, "%Y-%m-%d")
        source_start = datetime.strptime(start_date, "%Y-%m-%d")
        
        db = get_db()
        cur = db.cursor()
        
        # 获取本周所有课程
        cur.execute("""
            SELECT subject, teacher_id, student_ids, classroom, class_time, duration, class_date
            FROM course_schedule
            WHERE class_date BETWEEN %s AND %s
              AND (status IS NULL OR status != 'cancelled')
        """, (start_date, end_date))
        
        courses = cur.fetchall()
        copied_count = 0
        
        for course in courses:
            old_date = course[6]
            # 计算星期偏移
            old_date_obj = old_date
            if isinstance(old_date, str):
                old_date_obj = datetime.strptime(old_date, "%Y-%m-%d")
            day_offset = (old_date_obj - source_start).days
            new_date = target + timedelta(days=day_offset)
            new_date_str = new_date.strftime("%Y-%m-%d")
            
            cur.execute("""
                INSERT INTO course_schedule 
                (subject, teacher_id, student_ids, classroom, class_date, class_time, duration, status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, 'scheduled')
            """, (course[0], course[1], course[2], course[3], new_date_str, course[4], course[5]))
            copied_count += 1
        
        db.commit()
        cur.close()
        db.close()
        
        return jsonify({"code": 200, "msg": "复制成功", "count": copied_count})
    except Exception as e:
        print(f"复制周课程错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500





@app.route("/api/schedule/export/excel", methods=["POST"])
def export_schedule_excel():
    """导出周课表为Excel - 支持多课程"""
    try:
        data = request.json
        start_date = data.get('start_date')
        end_date = data.get('end_date')

        if not start_date or not end_date:
            return jsonify({"code": 400, "msg": "缺少日期参数"}), 400

        db = get_db()
        cur = db.cursor()

        # 获取周课表数据 - 确保获取所有课程
        cur.execute("""
            SELECT 
                s.id,
                s.class_date,
                EXTRACT(DOW FROM s.class_date) as weekday,
                s.class_time,
                s.subject,
                s.classroom,
                COALESCE(s.status, 'scheduled') as status,
                t.name as teacher_name,
                s.student_ids
            FROM course_schedule s
            LEFT JOIN teacher t ON s.teacher_id = t.id
            WHERE s.class_date BETWEEN %s AND %s
              AND (s.status IS NULL OR s.status != 'cancelled')
            ORDER BY s.class_date, s.class_time, s.id
        """, (start_date, end_date))

        data = cur.fetchall()

        # 获取学生名称映射
        cur.execute("SELECT id, name FROM student")
        students_map = {row[0]: row[1] for row in cur.fetchall()}

        cur.close()
        db.close()

        # 整理数据结构 - 支持多课程（列表形式）
        weekdays = ['周一', '周二', '周三', '周四', '周五', '周六', '周日']
        time_slots = set()
        schedule_data = {}

        print(f"查询到 {len(data)} 条课程记录")

        for row in data:
            weekday_idx = int(row[2])
            weekday_idx = 6 if weekday_idx == 0 else weekday_idx - 1
            weekday_name = weekdays[weekday_idx]
            class_time = row[3]
            time_slots.add(class_time)

            # 解析学生名称
            student_names = []
            if row[8]:
                student_ids = [int(x) for x in row[8].split(',') if x]
                student_names = [students_map.get(sid, '') for sid in student_ids if students_map.get(sid)]

            key = f"{weekday_name}_{class_time}"

            # 初始化为列表
            if key not in schedule_data:
                schedule_data[key] = []

            # 添加课程到列表
            schedule_data[key].append({
                "id": row[0],
                "subject": row[4] or '',
                "teacher": row[7] or '',
                "classroom": row[5] or '',
                "students": '、'.join(student_names) if student_names else '集体课',
                "status": '已完成' if row[6] == 'completed' else '待上课'
            })

        # 打印调试信息
        for key, courses in schedule_data.items():
            if len(courses) > 1:
                print(f"{key} 有 {len(courses)} 门课程")

        # 排序时间段
        def get_start_time(time_str):
            return time_str.split('-')[0] if time_str else '00:00'

        sorted_time_slots = sorted(list(time_slots), key=get_start_time)

        # 创建工作簿
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = f"周课表_{start_date}_至_{end_date}"

        # 设置样式
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        center_alignment = Alignment(horizontal="center", vertical="center")
        left_alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # 第一行：标题
        ws.merge_cells(f'A1:{chr(64 + len(weekdays) + 1)}1')
        title_cell = ws['A1']
        title_cell.value = f"课 程 表  ({start_date} ~ {end_date})"
        title_cell.font = Font(bold=True, size=16)
        title_cell.alignment = center_alignment

        # 第二行：表头
        headers = ['时间'] + weekdays
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=2, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_alignment
            cell.border = thin_border

        # 写入课程数据
        row_num = 3
        for time_slot in sorted_time_slots:
            # 时间列
            cell = ws.cell(row=row_num, column=1, value=time_slot)
            cell.alignment = center_alignment
            cell.border = thin_border

            # 记录当前行每个单元格的课程数量，用于设置行高
            max_courses_in_row = 0

            for col, weekday in enumerate(weekdays, 2):
                key = f"{weekday}_{time_slot}"
                if key in schedule_data and len(schedule_data[key]) > 0:
                    courses = schedule_data[key]
                    max_courses_in_row = max(max_courses_in_row, len(courses))

                    # 构建多课程文本
                    course_lines = []
                    for idx, course in enumerate(courses, 1):
                        course_lines.append(f"【{idx}】{course['subject']}")
                        course_lines.append(f"   教师：{course['teacher']}")
                        course_lines.append(f"   教室：{course['classroom']}")
                        course_lines.append(f"   学生：{course['students']}")
                        if idx < len(courses):
                            course_lines.append("")  # 课程间空行

                    cell_value = '\n'.join(course_lines)
                    cell = ws.cell(row=row_num, column=col, value=cell_value)
                    cell.alignment = left_alignment
                else:
                    cell = ws.cell(row=row_num, column=col, value="")
                    cell.alignment = center_alignment
                cell.border = thin_border

            # 根据课程数量设置行高
            if max_courses_in_row > 0:
                ws.row_dimensions[row_num].height = 30 + (max_courses_in_row - 1) * 35
            else:
                ws.row_dimensions[row_num].height = 25

            row_num += 1

        # 设置列宽
        for col in range(1, len(weekdays) + 2):
            col_letter = chr(64 + col)
            ws.column_dimensions[col_letter].width = 24

        # 保存到内存
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'周课表_{start_date}_至_{end_date}.xlsx'
        )
    except Exception as e:
        print(f"导出课表错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500


# ==================== 排课管理模块 - 明日课程 ====================
@app.route("/api/schedule/tomorrow", methods=["GET"])
def schedule_tomorrow():
    """获取明天的课程"""
    try:
        # 使用北京时间计算明天
        beijing_now = get_beijing_time()
        tomorrow = (beijing_now + timedelta(days=1)).strftime("%Y-%m-%d")
        print(f"北京时间: {beijing_now}, 明天: {tomorrow}")

        db = get_db()
        cur = db.cursor()

        cur.execute("""
            SELECT 
                cs.id,
                cs.class_time,
                cs.subject,
                cs.classroom,
                COALESCE(t.name, '待分配') as teacher_name
            FROM course_schedule cs
            LEFT JOIN teacher t ON cs.teacher_id = t.id
            WHERE cs.class_date = %s
              AND (cs.status IS NULL OR cs.status != 'cancelled')
            ORDER BY cs.class_time
        """, (tomorrow,))

        data = cur.fetchall()
        cur.close()
        db.close()

        result = []
        for row in data:
            result.append({
                "id": row[0],
                "class_time": row[1],
                "subject": row[2] or '',
                "classroom": row[3] or '',
                "teacher_name": row[4]
            })

        print(f"明天({tomorrow})有 {len(result)} 门课程")

        return jsonify({"code": 200, "data": result, "date": tomorrow})
    except Exception as e:
        print(f"获取明日课程错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


# 学生课时统计
@app.route("/api/student/attendance/statistics", methods=["GET"])
def get_student_attendance_statistics():
    """获取学生出勤统计"""
    try:
        student_id = request.args.get('student_id', type=int)
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        if not student_id:
            return jsonify({"code": 400, "msg": "请选择学生"}), 400

        db = get_db()
        cur = db.cursor()

        student_id_str = str(student_id)

        sql = """
            SELECT 
                cs.class_date,
                cs.class_time,
                cs.status
            FROM course_schedule cs
            WHERE cs.status != 'cancelled'
              AND (
                  cs.student_id = %s 
                  OR cs.student_ids = %s
                  OR cs.student_ids LIKE %s
                  OR cs.student_ids LIKE %s
                  OR cs.student_ids LIKE %s
              )
        """

        params = [
            student_id,
            student_id_str,
            f'{student_id_str},%',
            f'%,{student_id_str}',
            f'%,{student_id_str},%'
        ]

        if start_date:
            sql += " AND cs.class_date >= %s"
            params.append(start_date)
        if end_date:
            sql += " AND cs.class_date <= %s"
            params.append(end_date)

        cur.execute(sql, params)
        data = cur.fetchall()
        cur.close()
        db.close()

        now = datetime.now()
        today_date = now.date()
        now_time_str = now.strftime("%H:%M")

        total_classes = 0
        completed_classes = 0
        upcoming_classes = 0

        for row in data:
            class_date = row[0]
            class_time = row[1]
            status = row[2]

            if status == 'cancelled':
                continue

            total_classes += 1

            start_time_str = class_time.split('-')[0].strip() if class_time else "00:00"

            if class_date < today_date:
                completed_classes += 1
            elif class_date == today_date and start_time_str <= now_time_str:
                completed_classes += 1
            else:
                upcoming_classes += 1

        total_hours = total_classes * 2
        completed_hours = completed_classes * 2
        upcoming_hours = upcoming_classes * 2

        return jsonify({
            "code": 200,
            "data": {
                "total_classes": total_classes,
                "completed_classes": completed_classes,
                "upcoming_classes": upcoming_classes,
                "total_hours": total_hours,
                "completed_hours": completed_hours,
                "upcoming_hours": upcoming_hours
            }
        })
    except Exception as e:
        print(f"获取出勤统计错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


# ==================== 学生课时统计（基于排课表）====================
@app.route("/api/student/attendance/list", methods=["GET"])
def get_student_attendance_list():
    """获取学生的出勤记录列表"""
    try:
        student_id = request.args.get('student_id', type=int)
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        if not student_id:
            return jsonify({"code": 400, "msg": "请选择学生"}), 400

        db = get_db()
        cur = db.cursor()

        student_id_str = str(student_id)

        sql = """
            SELECT 
                cs.id,
                cs.class_date,
                cs.class_time,
                cs.subject,
                cs.classroom,
                cs.status,
                t.name as teacher_name
            FROM course_schedule cs
            LEFT JOIN teacher t ON cs.teacher_id = t.id
            WHERE cs.status != 'cancelled'
              AND (
                  cs.student_id = %s 
                  OR cs.student_ids = %s
                  OR cs.student_ids LIKE %s
                  OR cs.student_ids LIKE %s
                  OR cs.student_ids LIKE %s
              )
        """

        params = [
            student_id,
            student_id_str,
            f'{student_id_str},%',
            f'%,{student_id_str}',
            f'%,{student_id_str},%'
        ]

        if start_date:
            sql += " AND cs.class_date >= %s"
            params.append(start_date)
        if end_date:
            sql += " AND cs.class_date <= %s"
            params.append(end_date)

        sql += " ORDER BY cs.class_date DESC, cs.class_time"

        cur.execute(sql, params)
        data = cur.fetchall()
        cur.close()
        db.close()

        now = datetime.now()
        today_date = now.date()
        now_time_str = now.strftime("%H:%M")

        result = []
        for row in data:
            class_date = row[1]
            class_time = row[2]
            status = row[5]

            start_time_str = class_time.split('-')[0].strip() if class_time else "00:00"

            is_completed = False
            if class_date < today_date:
                is_completed = True
            elif class_date == today_date and start_time_str <= now_time_str:
                is_completed = True

            if is_completed and status != 'cancelled':
                actual_status = 'completed'
                status_text = '已上课'
            elif status == 'cancelled':
                actual_status = 'cancelled'
                status_text = '已取消'
            else:
                actual_status = 'scheduled'
                status_text = '待上课'

            result.append({
                "id": row[0],
                "class_date": str(class_date),
                "class_time": class_time,
                "subject": row[3] or '',
                "classroom": row[4] or '',
                "status": actual_status,
                "status_text": status_text,
                "teacher_name": row[6] or '待分配'
            })

        return jsonify({"code": 200, "data": result})
    except Exception as e:
        print(f"获取出勤记录错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


# ==================== 学生出勤报表导出接口 ====================
@app.route("/api/student/attendance/export", methods=["POST"])
def export_attendance_report():
    """导出学生出勤报表"""
    try:
        data = request.json
        student_ids = data.get('student_ids', [])
        start_date = data.get('start_date')
        end_date = data.get('end_date')

        # 1. 校验参数
        if not student_ids:
            return jsonify({"code": 400, "msg": "请选择学生"}), 400

        db = get_db()
        cur = db.cursor()
        all_records = []

        # 2. 循环查询每个学生的数据
        for student_id in student_ids:
            student_id_str = str(student_id)
            sql = """
                SELECT 
                    s.name as student_name,
                    s.grade,
                    s.phone,
                    cs.class_date,
                    cs.class_time,
                    cs.subject,
                    cs.classroom,
                    t.name as teacher_name,
                    cs.status
                FROM course_schedule cs
                LEFT JOIN teacher t ON cs.teacher_id = t.id
                CROSS JOIN student s
                WHERE s.id = %s
                  AND cs.status != 'cancelled'
                  AND (
                      cs.student_id = %s 
                      OR cs.student_ids = %s
                      OR cs.student_ids LIKE %s
                      OR cs.student_ids LIKE %s
                      OR cs.student_ids LIKE %s
                  )
            """
            params = [student_id, student_id, student_id_str,
                      f'{student_id_str},%', f'%,{student_id_str}', f'%,{student_id_str},%']

            if start_date:
                sql += " AND cs.class_date >= %s"
                params.append(start_date)
            if end_date:
                sql += " AND cs.class_date <= %s"
                params.append(end_date)

            sql += " ORDER BY cs.class_date, cs.class_time"
            cur.execute(sql, params)
            records = cur.fetchall()
            for row in records:
                all_records.append(row)

        cur.close()
        db.close()

        # 3. 生成 Excel 文件
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "学生出勤报表"
        headers = ["学生姓名", "年级", "联系电话", "上课日期", "上课时间", "课程名称", "教室", "授课教师", "状态"]
        ws.append(headers)

        # 设置样式
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        today_date = datetime.now().date()
        now_time_str = datetime.now().strftime("%H:%M")

        for row in all_records:
            class_date = row[3]
            class_time = row[4]
            start_time_str = class_time.split('-')[0].strip() if class_time else "00:00"

            if class_date < today_date:
                status_text = '已上课'
            elif class_date == today_date and start_time_str <= now_time_str:
                status_text = '已上课'
            else:
                status_text = '待上课'

            ws.append([
                row[0] or '', row[1] or '', row[2] or '',
                str(row[3]) if row[3] else '', row[4] or '',
                row[5] or '', row[6] or '', row[7] or '', status_text
            ])

        # 调整列宽
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws.column_dimensions[col_letter].width = min(max_length + 2, 25)

        # 4. 返回文件
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'学生出勤报表_{datetime.now().strftime("%Y%m%d")}.xlsx'
        )
    except Exception as e:
        print(f"导出出勤报表错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500


# ==================== 微信提醒模块 ====================

import requests
import json

# 微信配置（请确认这些值是否正确）
WECHAT_APP_ID = os.environ.get('WECHAT_APP_ID', 'wx7f3bff31a3dbfd0c')
WECHAT_APP_SECRET = os.environ.get('WECHAT_APP_SECRET', '74e6b9ccbf7495205aa5e1da0a30135e')


def get_access_token():
    """获取微信access_token"""
    try:
        url = f"https://api.weixin.qq.com/cgi-bin/token?grant_type=client_credential&appid={WECHAT_APP_ID}&secret={WECHAT_APP_SECRET}"
        response = requests.get(url, timeout=10)
        data = response.json()
        if 'access_token' in data:
            print("获取access_token成功")
            return data['access_token']
        else:
            print(f"获取access_token失败: {data}")
            return None
    except Exception as e:
        print(f"获取access_token错误: {str(e)}")
        return None


@app.route("/api/remind/send-tomorrow", methods=["POST"])
def send_tomorrow_remind():
    print("[接口] /api/remind/send-tomorrow 被调用")
    result = send_tomorrow_remind_internal()
    print(f"[接口] 返回结果: {result}")
    return jsonify(result)


@app.route("/api/remind/subscribe", methods=["POST"])
def subscribe_remind():
    """记录用户订阅状态"""
    try:
        data = request.json
        openid = data.get('openid')
        template_id = data.get('template_id')
        user_type = data.get('user_type', 'parent')

        # 临时：如果前端没传openid，从当前登录用户获取（需要根据你的用户系统实现）
        if not openid:
            # 这里简化处理，实际应该从session或token获取
            # 如果无法获取，可以暂时允许为空，只记录订阅
            return jsonify({"code": 200, "msg": "订阅记录已保存（openid缺失）"})

        db = get_db()
        cur = db.cursor()

        # 确保表存在
        cur.execute("""
            CREATE TABLE IF NOT EXISTS subscribe_record (
                id SERIAL PRIMARY KEY,
                openid VARCHAR(100) NOT NULL,
                template_id VARCHAR(100) NOT NULL,
                user_type VARCHAR(20),
                status VARCHAR(20) DEFAULT 'active',
                subscribe_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(openid, template_id)
            )
        """)

        cur.execute("""
            INSERT INTO subscribe_record (openid, template_id, user_type, status, subscribe_time)
            VALUES (%s, %s, %s, 'active', NOW())
            ON CONFLICT (openid, template_id) 
            DO UPDATE SET status = 'active', subscribe_time = NOW()
        """, (openid, template_id, user_type))

        db.commit()
        cur.close()
        db.close()

        return jsonify({"code": 200, "msg": "订阅成功"})
    except Exception as e:
        print(f"订阅错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/test/env", methods=["GET"])
def test_env():
    """测试环境变量"""
    return jsonify({
        "WECHAT_APP_ID": os.environ.get('WECHAT_APP_ID', 'NOT_SET'),
        "WECHAT_APP_SECRET": "SET" if os.environ.get('WECHAT_APP_SECRET') else 'NOT_SET',
        "DB_HOST": os.environ.get('DB_HOST', 'NOT_SET')
    })


# ==================== 定时任务模块 ====================

# from apscheduler.schedulers.background import BackgroundScheduler
# from apscheduler.triggers.cron import CronTrigger
# from datetime import datetime, timedelta

# 定时发送提醒函数
def scheduled_send_remind():
    """定时发送明天的课程提醒"""
    with app.app_context():
        print(f"[定时任务] 开始执行 - {datetime.now()}")
        try:
            # 调用发送提醒函数
            result = send_tomorrow_remind_internal()
            print(f"[定时任务] 执行结果: {result.get('msg', '未知')}")
        except Exception as e:
            print(f"[定时任务] 执行错误: {str(e)}")


def send_tomorrow_remind_internal():
    """内部发送明日课程提醒（给教师和家长）"""
    print("开始发送提醒，准备记录日志")
    try:
        beijing_now = datetime.utcnow() + timedelta(hours=8)
        tomorrow = (beijing_now + timedelta(days=1)).strftime("%Y-%m-%d")
        formatted_date = beijing_now.strftime("%Y年%m月%d日")
        
        db = get_db()
        cur = db.cursor()
        
        cur.execute("""
            SELECT 
                cs.id,
                cs.class_time,
                cs.subject,
                cs.classroom,
                cs.teacher_id,
                cs.student_ids
            FROM course_schedule cs
            WHERE cs.class_date = %s
              AND (cs.status IS NULL OR cs.status != 'cancelled')
        """, (tomorrow,))
        
        courses = cur.fetchall()
        
        if not courses:
            cur.close()
            db.close()
            return {"code": 200, "msg": "明天没有课程", "count": 0}
        
        access_token = get_access_token()
        if not access_token:
            cur.close()
            db.close()
            return {"code": 500, "msg": "获取access_token失败"}
        
        TEMPLATE_USER = "hEY6ukiBlTm79MQ4GL0heVpS0YDcHaiVWZAz3StSj0s"
        
        parent_sent = 0
        teacher_sent = 0
        
        for course in courses:
            course_id = course[0]
            class_time = course[1]
            subject = course[2] or '课程'
            teacher_id = course[4]
            student_ids_str = course[5] or ''
            
            start_time = class_time.split('-')[0] if class_time else "09:00"
            full_time_str = f"{formatted_date} {start_time}"
            
            # ========== 发送给教师 ==========
            if teacher_id:
                cur.execute("SELECT u.openid, u.id, u.name FROM \"user\" u WHERE u.teacher_id = %s AND u.openid IS NOT NULL", (teacher_id,))
                teacher_user = cur.fetchone()
                
                if teacher_user and teacher_user[0]:
                    send_data = {
                        "touser": teacher_user[0],
                        "template_id": TEMPLATE_USER,
                        "data": {
                            "date1": {"value": full_time_str},
                            "thing6": {"value": subject},
                            "short_thing20": {"value": "教师提醒"}
                        }
                    }
                    url = f"https://api.weixin.qq.com/cgi-bin/message/subscribe/send?access_token={access_token}"
                    response = requests.post(url, json=send_data, timeout=10)
                    result = response.json()
                    
                    is_success = (result.get('errcode') == 0)
                    
                    # 记录日志（一次性）
                    log_remind(
                        schedule_id=course_id,
                        remind_type='before_class',
                        receiver_role='teacher',
                        receiver_id=teacher_user[1],
                        receiver_openid=teacher_user[0],
                        receiver_phone=None,
                        content=f"课程提醒: {subject} {full_time_str}",
                        status='success' if is_success else 'failed',
                        error_msg=result.get('errmsg') if not is_success else None,
                        response_data=json.dumps(result)
                    )
                    
                    if is_success:
                        teacher_sent += 1
            
            # ========== 发送给家长 ==========
            if student_ids_str:
                student_ids = [int(x) for x in student_ids_str.split(',') if x]
                for sid in student_ids:
                    cur.execute("SELECT s.name, s.parent_phone FROM student s WHERE s.id = %s", (sid,))
                    student = cur.fetchone()
                    if student and student[1]:
                        cur.execute("SELECT u.openid, u.id, u.name FROM \"user\" u WHERE u.phone = %s AND u.openid IS NOT NULL", (student[1],))
                        parent_user = cur.fetchone()
                        
                        if parent_user and parent_user[0]:
                            send_data = {
                                "touser": parent_user[0],
                                "template_id": TEMPLATE_USER,
                                "data": {
                                    "date1": {"value": full_time_str},
                                    "thing6": {"value": subject},
                                    "short_thing20": {"value": student[0]}
                                }
                            }
                            url = f"https://api.weixin.qq.com/cgi-bin/message/subscribe/send?access_token={access_token}"
                            response = requests.post(url, json=send_data, timeout=10)
                            result = response.json()
                            
                            is_success = (result.get('errcode') == 0)
                            
                            # 记录日志（一次性）
                            log_remind(
                                schedule_id=course_id,
                                remind_type='before_class',
                                receiver_role='parent',
                                receiver_id=parent_user[1],
                                receiver_openid=parent_user[0],
                                receiver_phone=student[1],
                                content=f"学生{student[0]}的课程提醒: {subject} {full_time_str}",
                                status='success' if is_success else 'failed',
                                error_msg=result.get('errmsg') if not is_success else None,
                                response_data=json.dumps(result)
                            )
                            
                            if is_success:
                                parent_sent += 1
        
        cur.close()
        db.close()
        
        return {
            "code": 200,
            "msg": f"课前提醒发送完成：教师{teacher_sent}人，家长{parent_sent}人",
            "teacher_sent": teacher_sent,
            "parent_sent": parent_sent
        }
    except Exception as e:
        print(f"课前提醒发送错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"code": 500, "msg": str(e)}






def send_today_confirm_internal():
    """内部发送今日课程确认提醒（给管理员）"""
    try:
        db = get_db()
        cur = db.cursor()
        
        cur.execute("SELECT CURRENT_DATE")
        db_today = cur.fetchone()[0]
        today = db_today.strftime("%Y-%m-%d")
        formatted_date = db_today.strftime("%Y年%m月%d日")
        
        print(f"[课后确认] 查询日期: {today}")
        
        cur.execute("""
            SELECT 
                cs.id,
                cs.class_time,
                cs.subject,
                cs.classroom,
                cs.student_ids
            FROM course_schedule cs
            WHERE cs.class_date = %s
              AND (cs.status IS NULL OR cs.status = 'scheduled')
        """, (today,))
        
        courses = cur.fetchall()
        print(f"[课后确认] 找到 {len(courses)} 门待确认课程")
        
        if not courses:
            cur.close()
            db.close()
            return {"code": 200, "msg": "今天没有待确认的课程", "count": 0}
        
        access_token = get_access_token()
        if not access_token:
            cur.close()
            db.close()
            return {"code": 500, "msg": "获取access_token失败"}
        
        TEMPLATE_ADMIN = "qsPScuGxWPjB69boSJvaIleKJFSLJl-d6NRTLypPuYo"
        
        # 获取管理员openid
        cur.execute("SELECT u.openid, u.id, u.name FROM \"user\" u WHERE u.role = 'admin' AND u.openid IS NOT NULL LIMIT 1")
        admin_user = cur.fetchone()
        
        if not admin_user or not admin_user[0]:
            cur.close()
            db.close()
            return {"code": 500, "msg": "未找到管理员"}
        
        admin_openid = admin_user[0]
        admin_id = admin_user[1]
        confirm_count = 0
        
        for course in courses:
            course_id = course[0]
            class_time = course[1]
            subject = course[2] or '课程'
            student_ids_str = course[4] or ''
            
            start_time = class_time.split('-')[0] if class_time else "09:00"
            full_time_str = f"{formatted_date} {start_time}"
            
            # 获取学生名称
            student_names = []
            if student_ids_str:
                student_ids = [int(x) for x in student_ids_str.split(',') if x]
                for sid in student_ids:
                    cur.execute("SELECT name FROM student WHERE id = %s", (sid,))
                    student = cur.fetchone()
                    if student:
                        student_names.append(student[0])
            students_str = '、'.join(student_names) if student_names else '集体课'
            
            send_data = {
                "touser": admin_openid,
                "template_id": TEMPLATE_ADMIN,
                "data": {
                    "thing1": {"value": subject},
                    "time3": {"value": full_time_str},
                    "thing5": {"value": students_str + " 请确认是否上课！"}
                },
                "miniprogram": {
                    "appid": WECHAT_APP_ID,
                    "pagepath": f"pages/schedule/calendar/calendar?confirm_id={course_id}"
                }
            }
            
            url = f"https://api.weixin.qq.com/cgi-bin/message/subscribe/send?access_token={access_token}"
            response = requests.post(url, json=send_data, timeout=10)
            result = response.json()
            
            print(f"[课后确认] 课程 {subject} 发送结果: {result}")
            
            # 记录日志
            try:
                log_remind(
                    schedule_id=course_id,
                    remind_type='after_class',
                    receiver_role='admin',
                    receiver_id=admin_id,
                    receiver_openid=admin_openid,
                    receiver_phone=None,
                    content=f"课程确认提醒: {subject} {full_time_str} 学生: {students_str}",
                    status='success' if result.get('errcode') == 0 else 'failed',
                    error_msg=result.get('errmsg') if result.get('errcode') != 0 else None,
                    response_data=json.dumps(result)
                )
            except Exception as log_err:
                print(f"[课后确认] 记录日志失败: {log_err}")
            
            if result.get('errcode') == 0:
                confirm_count += 1
        
        cur.close()
        db.close()
        
        return {
            "code": 200,
            "msg": f"课后确认提醒发送完成：{confirm_count} 条",
            "count": confirm_count
        }
    except Exception as e:
        print(f"课后确认提醒发送错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"code": 500, "msg": str(e)}





# 启动定时任务
def start_scheduler():
    scheduler = BackgroundScheduler()

    # 每天早上 9:00 发送明日课程提醒
    scheduler.add_job(
        func=scheduled_send_remind,
        trigger=CronTrigger(hour=9, minute=0, timezone='Asia/Shanghai'),
        id='daily_remind',
        replace_existing=True
    )

    # 每天晚上 20:00 发送今日课程确认提醒
    scheduler.add_job(
        func=scheduled_send_confirm,
        trigger=CronTrigger(hour=20, minute=0, timezone='Asia/Shanghai'),
        id='daily_confirm',
        replace_existing=True
    )

    scheduler.start()
    print("定时任务已启动")



@app.route("/api/remind/send-today-confirm", methods=["POST"])
def send_today_confirm():
    """发送今日课程确认提醒（管理员）"""
    try:
        beijing_now = get_beijing_time()
        today = beijing_now.strftime("%Y-%m-%d")
        formatted_date = beijing_now.strftime("%Y年%m月%d日")

        db = get_db()
        cur = db.cursor()

        # 查询今天已经完成的课程（当前时间之后的课程不算）
        cur.execute("""
            SELECT 
                cs.id,
                cs.class_time,
                cs.subject,
                cs.classroom,
                cs.teacher_id,
                cs.student_ids,
                cs.status
            FROM course_schedule cs
            WHERE cs.class_date = %s
              AND (cs.status IS NULL OR cs.status = 'scheduled')
            ORDER BY cs.class_time
        """, (today,))

        courses = cur.fetchall()

        if not courses:
            cur.close()
            db.close()
            return jsonify({"code": 200, "msg": "今天没有待确认的课程", "count": 0})

        access_token = get_access_token()
        if not access_token:
            cur.close()
            db.close()
            return jsonify({"code": 500, "msg": "获取access_token失败"}), 500

        TEMPLATE_ADMIN = "qsPScuGxWPjB69boSJvaIleKJFSLJl-d6NRTLypPuYo"

        # 获取管理员openid
        cur.execute("SELECT openid FROM \"user\" WHERE role = 'admin' LIMIT 1")
        admin_user = cur.fetchone()
        admin_openid = admin_user[0] if admin_user else None

        if not admin_openid:
            cur.close()
            db.close()
            return jsonify({"code": 500, "msg": "未找到管理员"}), 500

        confirm_count = 0

        for course in courses:
            course_id = course[0]
            class_time = course[1]
            subject = course[2] or '课程'
            student_ids_str = course[5] or ''

            start_time = class_time.split('-')[0] if class_time else "09:00"
            full_time_str = f"{formatted_date} {start_time}"

            # 获取学生名称
            student_names = []
            if student_ids_str:
                student_ids = [int(x) for x in student_ids_str.split(',') if x]
                for sid in student_ids:
                    cur.execute("SELECT name FROM student WHERE id = %s", (sid,))
                    student = cur.fetchone()
                    if student:
                        student_names.append(student[0])
            students_str = '、'.join(student_names) if student_names else '集体课'

            # 构建确认链接（小程序路径）
            confirm_url = f"pages/schedule/calendar/calendar?confirm_id={course_id}"

            send_data = {
                "touser": admin_openid,
                "template_id": TEMPLATE_ADMIN,
                "data": {
                    "thing1": {"value": subject},
                    "time3": {"value": full_time_str},
                    "thing5": {"value": students_str+"请进系统确认！"}
                },
                "miniprogram": {
                    "appid": "wx7f3bff31a3dbfd0c",
                    "pagepath": confirm_url
                }
            }
            url = f"https://api.weixin.qq.com/cgi-bin/message/subscribe/send?access_token={access_token}"
            response = requests.post(url, json=send_data, timeout=10)
            result = response.json()
            if result.get('errcode') == 0:
                confirm_count += 1

        cur.close()
        db.close()

        return jsonify({
            "code": 200,
            "msg": f"已发送 {confirm_count} 条课程确认提醒",
            "count": confirm_count
        })
    except Exception as e:
        print(f"发送确认提醒错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500


# ==================== 课程确认功能 ====================

@app.route("/api/schedule/confirm", methods=["POST"])
def confirm_schedule():
    """管理员确认课程完成，并自动记录课时消耗"""
    try:
        data = request.json
        course_id = data.get('course_id')
        
        if not course_id:
            return jsonify({"code": 400, "msg": "缺少课程ID"}), 400
        
        db = get_db()
        cur = db.cursor()
        
        # 1. 检查课程是否存在以及是否已确认
        cur.execute("SELECT id, subject, status, duration, student_id, student_ids FROM course_schedule WHERE id = %s", (course_id,))
        course = cur.fetchone()
        
        if not course:
            cur.close()
            db.close()
            return jsonify({"code": 404, "msg": "课程不存在"}), 404
        
        if course[2] == 'completed':
            cur.close()
            db.close()
            return jsonify({"code": 400, "msg": "课程已完成确认，不能重复确认"}), 400
        
        # 2. 检查是否已经有课时消耗记录（防止重复扣减）
        cur.execute("SELECT COUNT(*) FROM hour_consumption WHERE schedule_id = %s", (course_id,))
        existing_count = cur.fetchone()[0]
        if existing_count > 0:
            # 已经有记录，说明可能重复调用，但课程状态还是 scheduled
            # 这种情况应该直接更新课程状态，不再重复扣减
            cur.execute("UPDATE course_schedule SET status = 'completed' WHERE id = %s", (course_id,))
            db.commit()
            cur.close()
            db.close()
            return jsonify({"code": 200, "msg": "课程已确认（消耗记录已存在）"}), 200
        
        # 3. 更新课程状态
        cur.execute("UPDATE course_schedule SET status = 'completed' WHERE id = %s", (course_id,))
        
        # 4. 获取课程涉及的课时数
        duration = float(course[3]) if course[3] else 2
        
        # 5. 获取学生ID列表
        student_ids = []
        if course[4]:  # student_id（单学生）
            student_ids.append(course[4])
        if course[5]:  # student_ids（多学生，逗号分隔）
            for sid in course[5].split(','):
                if sid and int(sid) not in student_ids:
                    student_ids.append(int(sid))
        
        # 6. 为每个学生记录课时消耗
        consumption_count = 0
        for student_id in student_ids:
            # 查找该学生有剩余课时的课时包
            cur.execute("""
                SELECT id, surplus FROM course_package 
                WHERE student_id = %s 
                  AND status = 'active' 
                  AND surplus > 0
                ORDER BY created_at ASC
                LIMIT 1
            """, (student_id,))
            
            package = cur.fetchone()
            if package:
                package_id = package[0]
                surplus = float(package[1]) if package[1] else 0
                
                # 消耗课时（优先从最早的课时包扣减）
                consume_hours = min(duration, surplus)
                
                # 更新课时包
                cur.execute("""
                    UPDATE course_package 
                    SET used = used + %s, surplus = surplus - %s 
                    WHERE id = %s
                """, (consume_hours, consume_hours, package_id))
                
                # 记录课时消耗（添加唯一约束防止重复）
                cur.execute("""
                    INSERT INTO hour_consumption (package_id, schedule_id, hours, consume_date, note)
                    VALUES (%s, %s, %s, CURRENT_DATE, %s)
                """, (package_id, course_id, consume_hours, f"课程确认：{course[1]}"))
                
                consumption_count += 1
        
        db.commit()
        cur.close()
        db.close()
        
        return jsonify({
            "code": 200,
            "msg": "确认成功，已记录课时消耗",
            "data": {
                "course_id": course_id,
                "subject": course[1],
                "consumption_count": consumption_count
            }
        })
    except Exception as e:
        print(f"确认课程错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/schedule/confirm/batch", methods=["POST"])
def batch_confirm_schedule():
    """批量确认课程完成，并自动记录课时消耗"""
    try:
        data = request.json
        course_ids = data.get('course_ids', [])
        
        if not course_ids:
            return jsonify({"code": 400, "msg": "请选择课程"}), 400
        
        db = get_db()
        cur = db.cursor()
        
        # 查询待确认课程
        placeholders = ','.join(['%s'] * len(course_ids))
        cur.execute(f"""
            SELECT id, subject, status, duration, student_id, student_ids 
            FROM course_schedule 
            WHERE id IN ({placeholders})
        """, course_ids)
        
        courses = cur.fetchall()
        
        confirmed_count = 0
        for course in courses:
            if course[2] == 'completed':
                continue
            
            # 更新课程状态
            cur.execute("UPDATE course_schedule SET status = 'completed' WHERE id = %s", (course[0],))
            
            duration = float(course[3]) if course[3] else 2
            
            # 获取学生ID列表
            student_ids = []
            if course[4]:
                student_ids.append(course[4])
            if course[5]:
                for sid in course[5].split(','):
                    if sid and int(sid) not in student_ids:
                        student_ids.append(int(sid))
            
            # 记录课时消耗
            for student_id in student_ids:
                cur.execute("""
                    SELECT id, surplus FROM course_package 
                    WHERE student_id = %s AND status = 'active' AND surplus > 0
                    ORDER BY created_at ASC
                    LIMIT 1
                """, (student_id,))
                
                package = cur.fetchone()
                if package:
                    consume_hours = min(duration, float(package[1]))
                    cur.execute("""
                        UPDATE course_package SET used = used + %s, surplus = surplus - %s WHERE id = %s
                    """, (consume_hours, consume_hours, package[0]))
                    
                    cur.execute("""
                        INSERT INTO hour_consumption (package_id, schedule_id, hours, consume_date, note)
                        VALUES (%s, %s, %s, CURRENT_DATE, %s)
                    """, (package[0], course[0], consume_hours, f"批量确认：{course[1]}"))
            
            confirmed_count += 1
        
        db.commit()
        cur.close()
        db.close()
        
        return jsonify({
            "code": 200,
            "msg": f"成功确认 {confirmed_count} 门课程",
            "count": confirmed_count
        })
    except Exception as e:
        print(f"批量确认错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500

@app.route("/api/student/hours/statistics-v2", methods=["GET"])


@app.route("/api/schedule/list", methods=["GET"])
def schedule_list():
    """获取排课列表（支持日期筛选）"""
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        db = get_db()
        cur = db.cursor()
        
        sql = """
            SELECT 
                cs.id,
                cs.class_date,
                cs.class_time,
                cs.subject,
                cs.classroom,
                cs.status,
                t.name as teacher_name,
                cs.student_ids
            FROM course_schedule cs
            LEFT JOIN teacher t ON cs.teacher_id = t.id
            WHERE 1=1
        """
        params = []
        
        if start_date:
            sql += " AND cs.class_date >= %s"
            params.append(start_date)
        if end_date:
            sql += " AND cs.class_date <= %s"
            params.append(end_date)
        
        sql += " ORDER BY cs.class_date DESC, cs.class_time"
        
        cur.execute(sql, params)
        rows = cur.fetchall()
        
        # 获取学生名称映射
        cur.execute("SELECT id, name FROM student")
        student_map = {row[0]: row[1] for row in cur.fetchall()}
        
        cur.close()
        db.close()
        
        result = []
        for row in rows:
            student_ids_str = row[7] or ''
            student_ids = [int(x) for x in student_ids_str.split(',') if x]
            student_names = [student_map.get(sid, '') for sid in student_ids if student_map.get(sid)]
            result.append({
                "id": row[0],
                "class_date": str(row[1]),
                "class_time": row[2],
                "subject": row[3] or '',
                "classroom": row[4] or '',
                "status": row[5] or 'scheduled',
                "teacher_name": row[6] or '待分配',
                "students": student_names,
                "student_count": len(student_names)
            })
        
        return jsonify({"code": 200, "data": result})
    except Exception as e:
        print(f"获取排课列表错误: {str(e)}")
        return jsonify({"code": 500, "msg": str(e)}), 500

# @app.route("/api/schedule/auto-confirm", methods=["POST"])
# def auto_confirm_courses():
#     """自动确认今天及以前已结束的课程"""
#     try:
#         db = get_db()
#         cur = db.cursor()

#         # 获取北京时间
#         from datetime import datetime, timedelta
#         beijing_now = datetime.utcnow() + timedelta(hours=8)
#         today = beijing_now.strftime("%Y-%m-%d")
#         current_time = beijing_now.strftime("%H:%M")

#         # 确认今天开始时间已过的课程
#         cur.execute("""
#             UPDATE course_schedule 
#             SET status = 'completed' 
#             WHERE class_date <= %s
#               AND (class_date < %s OR class_time <= %s)
#               AND (status IS NULL OR status != 'completed')
#               AND status != 'cancelled'
#         """, (today, today, current_time))

#         updated_count = cur.rowcount
#         db.commit()
#         cur.close()
#         db.close()

#         return jsonify({
#             "code": 200,
#             "msg": f"成功确认 {updated_count} 门课程",
#             "count": updated_count
#         })
#     except Exception as e:
#         print(f"自动确认错误: {str(e)}")
#         return jsonify({"code": 500, "msg": str(e)}), 500


# ==================== 仪表盘数据 ====================
@app.route("/api/dashboard/stats")
def dashboard_stats():
    """获取首页统计数据"""
    try:
        db = get_db()
        cur = db.cursor()

        # 学生总数
        try:
            cur.execute("SELECT COUNT(*) FROM student WHERE status = 1")
            student_count = cur.fetchone()[0]
        except Exception:
            cur.execute("SELECT COUNT(*) FROM student")
            student_count = cur.fetchone()[0]

        # 教师总数
        try:
            cur.execute("SELECT COUNT(*) FROM teacher WHERE status = 'active'")
            teacher_count = cur.fetchone()[0]
        except Exception:
            cur.execute("SELECT COUNT(*) FROM teacher")
            teacher_count = cur.fetchone()[0]

        # 今日课程数
        today = datetime.now().strftime("%Y-%m-%d")
        cur.execute("""
            SELECT COUNT(*) FROM course_schedule 
            WHERE class_date = %s 
              AND (status IS NULL OR status != 'cancelled')
        """, (today,))
        today_classes = cur.fetchone()[0]

        # 待确认课程数 (所有未完成且未取消的课程，日期<=今天)
        cur.execute("""
            SELECT COUNT(*) FROM course_schedule 
            WHERE class_date <= %s
              AND (status IS NULL OR status = 'scheduled')
              AND status != 'cancelled'
        """, (today,))
        pending_classes = cur.fetchone()[0]

        cur.close()
        db.close()

        return jsonify({"code": 200, "data": {
            "student_count": student_count,
            "teacher_count": teacher_count,
            "today_classes": today_classes,
            "pending_classes": pending_classes   # 替换原来的 total_surplus_hours
        }})
    except Exception as e:
        print(f"仪表盘错误: {str(e)}")
        return jsonify({"code": 200, "data": {
            "student_count": 0,
            "teacher_count": 0,
            "today_classes": 0,
            "pending_classes": 0
        }})


# ==================== 搜索接口 ====================
@app.route("/api/search/teachers", methods=["GET"])
def search_teachers():
    """模糊查询教师"""
    try:
        keyword = request.args.get('keyword', '')
        db = get_db()
        cur = db.cursor()
        cur.execute("""
            SELECT id, name, phone, subject 
            FROM teacher 
            WHERE status = 'active' 
            AND (name LIKE %s OR subject LIKE %s)
            LIMIT 20
        """, (f'%{keyword}%', f'%{keyword}%'))
        data = cur.fetchall()
        cur.close()
        db.close()

        result = [{"id": r[0], "name": r[1], "phone": r[2], "subject": r[3]} for r in data]
        return jsonify({"code": 200, "data": result})
    except Exception as e:
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/search/students", methods=["GET"])
def search_students():
    """模糊查询学生"""
    try:
        keyword = request.args.get('keyword', '')
        db = get_db()
        cur = db.cursor()
        cur.execute("""
            SELECT id, name, phone, grade 
            FROM student 
            WHERE status = 1 
            AND (name LIKE %s OR phone LIKE %s)
            LIMIT 30
        """, (f'%{keyword}%', f'%{keyword}%'))
        data = cur.fetchall()
        cur.close()
        db.close()

        result = [{"id": r[0], "name": r[1], "phone": r[2], "grade": r[3]} for r in data]
        return jsonify({"code": 200, "data": result})
    except Exception as e:
        return jsonify({"code": 500, "msg": str(e)}), 500


@app.route("/api/courses/list", methods=["GET"])
def get_courses_list():
    """获取课程列表"""
    try:
        db = get_db()
        cur = db.cursor()
        cur.execute("""
            SELECT DISTINCT subject FROM course_schedule WHERE subject IS NOT NULL AND subject != ''
            UNION
            SELECT '数学' as subject
            UNION
            SELECT '语文'
            UNION
            SELECT '英语'
            UNION
            SELECT '物理'
            UNION
            SELECT '化学'
            ORDER BY subject
        """)
        data = cur.fetchall()
        cur.close()
        db.close()

        result = [r[0] for r in data if r[0]]
        return jsonify({"code": 200, "data": result})
    except Exception as e:
        return jsonify({"code": 200, "data": ['数学', '语文', '英语', '物理', '化学']}), 200


@app.route("/api/rooms/list", methods=["GET"])
def get_rooms_list():
    """获取教室列表"""
    try:
        db = get_db()
        cur = db.cursor()
        cur.execute("""
            SELECT DISTINCT classroom FROM course_schedule WHERE classroom IS NOT NULL AND classroom != ''
            UNION
            SELECT 'A101' as classroom
            UNION
            SELECT 'A102'
            UNION
            SELECT 'A103'
            UNION
            SELECT 'B201'
            UNION
            SELECT 'B202'
            UNION
            SELECT 'C301'
            ORDER BY classroom
        """)
        data = cur.fetchall()
        cur.close()
        db.close()

        result = [r[0] for r in data if r[0]]
        return jsonify({"code": 200, "data": result})
    except Exception as e:
        return jsonify({"code": 200, "data": ['A101', 'A102', 'A103', 'B201', 'B202', 'C301']}), 200


@app.route("/api/teacher/salary/export", methods=["POST"])
def teacher_salary_export():
    """导出教师课酬报表"""
    try:
        data = request.json
        teacher_ids = data.get('teacher_ids', [])
        month = data.get('month')

        db = get_db()
        cur = db.cursor()

        # 获取课酬数据
        if teacher_ids:
            placeholders = ','.join(['%s'] * len(teacher_ids))
            sql = f"""
                SELECT 
                    t.name,
                    t.phone,
                    t.subject,
                    t.class_fee,
                    COUNT(CASE WHEN s.status = 'completed' THEN 1 END) as completed_classes,
                    COALESCE(SUM(CASE WHEN s.status = 'completed' THEN s.duration ELSE 0 END), 0) as total_hours,
                    COALESCE(SUM(CASE WHEN s.status = 'completed' THEN s.duration * t.class_fee ELSE 0 END), 0) as total_amount
                FROM teacher t
                LEFT JOIN course_schedule s ON t.id = s.teacher_id
                WHERE t.id IN ({placeholders})
                GROUP BY t.id, t.name, t.phone, t.subject, t.class_fee
            """
            cur.execute(sql, teacher_ids)
        else:
            cur.execute("""
                SELECT 
                    name, phone, subject, class_fee,
                    0 as completed_classes, 0 as total_hours, 0 as total_amount
                FROM teacher
                WHERE status = 'active'
            """)

        results = cur.fetchall()
        cur.close()
        db.close()

        # 创建Excel文件
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = f"教师课酬统计_{month or '全部'}"

        # 设置表头
        headers = ["教师姓名", "手机号", "教学科目", "课时费(元/小时)", "完成课时数", "授课总时长(小时)",
                   "应发课酬(元)"]
        ws.append(headers)

        # 设置表头样式
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        # 写入数据
        for row in results:
            ws.append([
                row[0] or '',
                row[1] or '',
                row[2] or '',
                float(row[3]) if row[3] else 0,
                row[4] or 0,
                float(row[5]) if row[5] else 0,
                float(row[6]) if row[6] else 0
            ])

        # 调整列宽
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 25)
            ws.column_dimensions[col_letter].width = adjusted_width

        # 保存到内存
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'教师课酬统计_{month or "全部"}.xlsx'
        )
    except Exception as e:
        print(f"导出课酬报表错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500


# ==================== 上课记录导出接口 ====================
@app.route("/api/schedule/export/attendance", methods=["POST"])
def export_attendance_record():
    """导出上课记录"""
    try:
        data = request.json
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        month = data.get('month')

        if not start_date or not end_date:
            return jsonify({"code": 400, "msg": "缺少日期参数"}), 400

        db = get_db()
        cur = db.cursor()

        # 查询该月份的所有已确认课程（兼容 PostgreSQL，不使用 GROUP_CONCAT）
        cur.execute("""
            SELECT 
                cs.class_date,
                cs.class_time,
                cs.subject,
                cs.classroom,
                t.name as teacher_name,
                cs.student_ids
            FROM course_schedule cs
            LEFT JOIN teacher t ON cs.teacher_id = t.id
            WHERE cs.class_date >= %s
              AND cs.class_date < %s
              AND cs.status = 'completed'
            ORDER BY cs.class_date, cs.class_time
        """, (start_date, end_date))

        courses = cur.fetchall()

        # 获取学生名称映射
        cur.execute("SELECT id, name FROM student")
        students_map = {row[0]: row[1] for row in cur.fetchall()}

        cur.close()
        db.close()

        # 创建Excel文件
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = f"上课记录_{month or '汇总'}"

        # 设置表头
        headers = ["上课日期", "上课时间", "课程名称", "教室", "授课教师", "上课学生"]
        ws.append(headers)

        # 设置表头样式
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        # 写入数据
        for course in courses:
            # 解析学生名称
            student_names = []
            if course[5]:  # student_ids
                student_ids = [int(x) for x in course[5].split(',') if x]
                student_names = [students_map.get(sid, '') for sid in student_ids if students_map.get(sid)]
            students_str = '、'.join(student_names) if student_names else '集体课'

            ws.append([
                str(course[0]) if course[0] else '',
                course[1] or '',
                course[2] or '',
                course[3] or '',
                course[4] or '待分配',
                students_str
            ])

        # 调整列宽
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 20)
            ws.column_dimensions[col_letter].width = adjusted_width

        # 保存并返回文件
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'上课记录_{month or "汇总"}.xlsx'
        )
    except Exception as e:
        print(f"导出上课记录错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "msg": str(e)}), 500


# ------------------- 启动应用 -------------------
if __name__ == "__main__":
     # 启动定时任务
    start_scheduler()
    port = int(os.getenv("PORT", 8080))
    app.run(host="0.0.0.0", port=port, debug=False)
   
